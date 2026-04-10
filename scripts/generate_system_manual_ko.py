# -*- coding: utf-8 -*-
from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DESKTOP = Path.home() / "Desktop"


def set_korean_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    style.font.size = Pt(10.5)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbers(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    for idx, line in enumerate(lines):
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9)
        if idx < len(lines) - 1:
            p.add_run("\n")


def build_manual(docx_path: Path, image_path: Path | None = None) -> None:
    doc = Document()
    set_korean_font(doc)

    doc.add_heading("PIPENET 결과물 검증 서버 상세 설명서", 0)
    meta = doc.add_paragraph()
    meta.add_run(
        f"문서 버전: 2026-04-10\n"
        f"작성일: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n"
        f"대상 프로젝트: pipenet_server_source_2026-04-10\n"
        f"대상 경로: {PROJECT_ROOT}"
    )
    doc.add_paragraph(
        "본 문서는 현재 PyCharm 프로젝트의 실제 서버 구조를 기준으로 작성했으며, "
        "백엔드/프론트엔드/검증 로직/운영 프로세스를 구체적으로 설명한다."
    )

    doc.add_heading("1. 문서 목적", level=1)
    doc.add_paragraph(
        "이 문서는 PIPENET 수리계산 검증 서버의 전체 구조를 개발자와 실무 사용자가 함께 이해할 수 있도록 정리한 기술 설명서이다. "
        "주요 목표는 다음과 같다."
    )
    add_bullets(
        doc,
        [
            "입력 파일이 어떤 경로로 서버에 들어오고 어떤 로직을 거쳐 검증되는지 설명",
            "백엔드와 프론트엔드의 역할 분담을 명확히 정리",
            "현재 구현된 규칙 엔진과 결과 데이터 테이블의 의미를 문서화",
            "운영, 배포, 장애 대응 시 참고 가능한 실무 문서로 활용",
        ],
    )

    doc.add_heading("2. 시스템 개요", level=1)
    doc.add_paragraph(
        "본 시스템은 Flask 기반 웹 서버이며, 사용자가 결과서(DOCX/PDF)와 SDF 파일을 업로드하면 "
        "서버가 결과서의 표 데이터를 구조화하고, 검증 기준에 따라 PASS/FAIL/REVIEW를 판정하여 "
        "웹 화면과 엑셀 파일로 제공한다."
    )
    add_bullets(
        doc,
        [
            "입력: 결과서 파일(docx/pdf), SDF 파일(선택), CAD 파일(선택)",
            "핵심 엔진: pipenet_validator.py",
            "웹 서버: 대조 서버.py",
            "화면 렌더링: templates/index.html + static/app.js + static/styles.css",
            "보조 모듈: cad_engine.py, head_detector.py",
        ],
    )

    doc.add_heading("3. 전체 데이터 흐름", level=1)
    doc.add_paragraph("시스템의 대표 데이터 흐름은 아래 순서로 진행된다.")
    add_numbers(
        doc,
        [
            "사용자가 메인 화면에서 결과서 파일과 SDF 파일을 업로드한다.",
            "브라우저는 /api/validate 로 FormData를 전송한다.",
            "대조 서버.py 가 업로드 파일을 data/uploads 폴더에 저장한다.",
            "PipenetGuideValidator 가 결과서를 읽고 섹션별 파서를 실행한다.",
            "배관, 헤드, 특수설비, 감압밸브, 공통 HW 검산, 배관 topology, 배관 규칙 엔진이 순차 실행된다.",
            "검증 결과는 summary/results/insights/stats/tables/rules/report 구조의 JSON으로 생성된다.",
            "대조 서버.py 는 추가로 시각화 그래프와 SDF 아이소매트릭 그래프 데이터를 붙여 응답한다.",
            "프론트엔드가 이 응답을 받아 검증결과/결과 데이터 테이블/통계/상세리포트 화면에 렌더링한다.",
        ],
    )
    doc.add_paragraph("참고용 간략 흐름도:")
    add_code_block(
        doc,
        [
            "사용자 업로드",
            "  -> /api/validate",
            "  -> 대조 서버.py",
            "  -> pipenet_validator.py",
            "  -> rules / stats / tables 생성",
            "  -> 시각화 + SDF graph 추가",
            "  -> JSON 응답",
            "  -> index.html / app.js 렌더링",
        ],
    )

    doc.add_heading("4. 백엔드 구조", level=1)
    doc.add_paragraph("백엔드는 Flask 서버와 검증 엔진, CAD 보조 모듈로 구성된다.")

    doc.add_heading("4.1 Flask 서버", level=2)
    doc.add_paragraph("핵심 파일: 대조 서버.py")
    add_bullets(
        doc,
        [
            "메인 화면 렌더링(/)",
            "검증 요청 처리(/api/validate)",
            "업데이트 기록 조회(/api/update-history)",
            "엑셀 내보내기(/api/export-xlsx)",
            "CAD 파싱 및 CAD 대조 API",
            "Matplotlib 기반 시각화 생성",
            "SDF 그래프 JSON 생성",
        ],
    )

    doc.add_heading("4.2 검증 엔진", level=2)
    doc.add_paragraph("핵심 파일: pipenet_validator.py")
    add_bullets(
        doc,
        [
            "결과서 텍스트 추출(DOCX XML / PDF text)",
            "섹션별 파서 실행",
            "공통 Hazen-Williams 선언 및 재계산 검증",
            "배관 topology 기반 유속 판정",
            "배관 대제목 규칙 엔진(PIPE.001~PIPE.006)",
            "헤드, 특수설비, 감압밸브 검증",
            "결과 테이블, 통계, structured rule result 생성",
        ],
    )

    doc.add_heading("4.3 CAD 보조 모듈", level=2)
    doc.add_paragraph("핵심 파일: cad_engine.py, head_detector.py")
    add_bullets(
        doc,
        [
            "DXF 읽기 및 네트워크 엔티티 직렬화",
            "레이어 기반 필터링 및 표시 대상 추출",
            "CAD-아이소매트릭 대조 모듈 지원",
            "명시적 mismatch 목록 기반 CAD 대조 helper 제공",
            "헤드 탐지(템플릿/YOLO) 보조",
        ],
    )

    doc.add_heading("5. 프론트엔드 구조", level=1)
    doc.add_paragraph("프론트엔드는 정적 HTML/CSS/JS로 구성되어 있으며 서버 응답을 메뉴형 인터페이스로 표현한다.")

    doc.add_heading("5.1 화면 골격", level=2)
    doc.add_paragraph("핵심 파일: templates/index.html")
    add_bullets(
        doc,
        [
            "좌측 메뉴: 검증결과 / 설계 최적화 가이드 / 결과 데이터 테이블 / 검진 통계 / 상세리포트 / CAD 대조 모듈",
            "상단 버튼: 업데이트 기록, 평가 기준",
            "중앙 업로드 폼: 결과서 파일(docx), SDF 파일(선택)",
            "각 패널은 필요한 경우에만 표시되는 구조",
        ],
    )

    doc.add_heading("5.2 동작 스크립트", level=2)
    doc.add_paragraph("핵심 파일: static/app.js")
    add_bullets(
        doc,
        [
            "업로드 폼 submit 처리",
            "검증 결과 테이블 렌더링",
            "결과 데이터 테이블 탭 전환",
            "필터 판정 로직 4칸 카드 모달",
            "배관 규칙 상태 보기 모달",
            "SDF 그래프 표시 및 줌/이동",
            "업데이트 기록/평가 기준 모달 제어",
        ],
    )

    doc.add_heading("5.3 스타일", level=2)
    doc.add_paragraph("핵심 파일: static/styles.css")
    add_bullets(
        doc,
        [
            "좌측 메뉴 + 우측 작업영역 레이아웃",
            "결과 카드, 테이블, 모달, 그래프 카드 스타일",
            "배관 테이블과 HW 상세 열 분리 배치",
            "반응형 규칙 및 모바일 축소 레이아웃",
        ],
    )

    doc.add_heading("6. 결과서 파싱 구조", level=1)
    doc.add_paragraph("검증 엔진은 결과서에서 아래 주요 섹션을 읽는다.")
    add_bullets(
        doc,
        [
            "DESIGN INFORMATION",
            "PIPE TYPES / AVAILABLE PIPE SIZES",
            "PIPE CONFIGURATION",
            "DESIGNED DIAMETERS & FLOWRATES",
            "NOZZLE CONFIGURATION",
            "FLOW IN PIPES",
            "FLOW THROUGH NOZZLES",
            "SPECIAL EQUIPMENT",
            "FLOW THROUGH ELASTOMERIC VALVES",
        ],
    )
    doc.add_paragraph("숫자 파싱은 공통 패턴으로 통일되어 scientific notation, trailing decimal, signed number를 처리한다.")
    add_code_block(doc, [r'NUM = r"[-+]?(?:\d+(?:\.\d*)?|\.\d+)(?:E[-+]?\d+)?"'])

    doc.add_heading("7. 현재 구현된 핵심 검증 항목", level=1)

    doc.add_heading("7.1 공통 규칙", level=2)
    add_bullets(
        doc,
        [
            "Hazen-Williams 선언 문구 존재 여부 확인",
            "Frict. Loss를 공식으로 재계산해 결과서 값과 비교",
        ],
    )

    doc.add_heading("7.2 배관 유속 규칙", level=2)
    doc.add_paragraph(
        "배관 유속은 더 이상 구경만으로 branch/main을 나누지 않고 topology 기반으로 판정한다."
    )
    add_bullets(
        doc,
        [
            "PIPE CONFIGURATION + NOZZLE CONFIGURATION으로 directed graph 구성",
            "downstream nozzle count 계산",
            "subtree cross split 여부 판단",
            "branch -> 6.0 m/s, other -> 10.0 m/s 적용",
            "ambiguity 발생 시 review 처리",
        ],
    )

    doc.add_heading("7.3 배관 대제목 규칙 엔진", level=2)
    add_bullets(
        doc,
        [
            "PIPE.001 도면과 schematic 일치 여부",
            "PIPE.002 1.2MPa 이상 구간 KSD3562 사용 여부",
            "PIPE.003 재질별 C factor 일치 여부",
            "PIPE.004 단위세대 내부 배관의 CPVC 사용 여부",
            "PIPE.005 단위세대 유입 배관의 65A 이하 여부",
            "PIPE.006 헤드 개수별 배관 구경 정책 충족 여부",
        ],
    )
    doc.add_paragraph(
        "이 규칙들은 results의 요약 문장뿐 아니라 rules.pipe[] structured result로도 내려가며, "
        "프론트에서는 배관 규칙 상태 보기 버튼으로 PASS/REVIEW/FAIL을 한 번에 확인할 수 있다."
    )

    doc.add_heading("7.4 헤드/특수설비/감압밸브", level=2)
    add_bullets(
        doc,
        [
            "헤드 최소 유량, 압력 범위 확인",
            "FX, AV, PV 등가길이 확인",
            "감압밸브 압력강하 계산 일치 여부 확인",
        ],
    )

    doc.add_heading("8. 결과 데이터 테이블 구조", level=1)
    doc.add_paragraph("배관 테이블은 현재 가장 많은 정보를 담고 있으며, 기본 표와 HW 상세 열로 나뉜다.")
    add_bullets(
        doc,
        [
            "기본 표: Pipe, 입력/출력 노드, 구경, 재질, 최대 관압, 배관 역할, 유속 기준, 유속 적합 등",
            "상세 열: 실내경, C-Factor, 길이, 등가길이, HW 재계산 마찰손실, 절대/상대오차 등",
            "structured rule 상태: 고압재질, C값, 세대내 CPVC, 세대유입 65A, 헤드수-구경정책",
        ],
    )
    doc.add_paragraph("현재 행 색상 정책은 다음과 같다.")
    add_bullets(
        doc,
        [
            "빨강: 실제 FAIL",
            "파랑: 공학 최적화 후보",
            "초록: 경제성 검토 후보",
            "REVIEW는 행 전체 노란색으로 칠하지 않고, 규칙 상태 값과 별도 버튼 모달로 확인",
        ],
    )

    doc.add_heading("9. 배관 규칙 상태 보기 기능", level=1)
    doc.add_paragraph(
        "검증 결과 패널에는 '배관 규칙 상태 보기' 버튼이 있으며, 배관 대제목 규칙의 메타 결과를 요약해서 보여준다."
    )
    add_bullets(
        doc,
        [
            "예: PIPE.002 고압구간 재질 확인: PASS",
            "예: PIPE.003 C값 확인: PASS",
            "예: PIPE.001 도면-스케매틱 일치: REVIEW",
            "예: PIPE.004 세대내 CPVC: REVIEW",
            "예: PIPE.005 세대유입 65A: REVIEW",
            "예: PIPE.006 헤드개수별 배관구경 정책: REVIEW",
        ],
    )

    doc.add_heading("10. 통계 및 시각화", level=1)
    add_bullets(
        doc,
        [
            "Pipe Velocity Check 그래프",
            "Nozzle Pressure-Flow 그래프",
            "결과서 통계 표",
            "SDF 통계 표",
            "배관 규칙 통계: 최대 관압, 고압 배관 수, C값 불일치 수, pipe review 수 등",
        ],
    )

    doc.add_heading("11. CAD-아이소매트릭 대조 모듈", level=1)
    doc.add_paragraph(
        "6번 메뉴는 독립형 듀얼 뷰어 모듈이며, DXF와 SDF를 나란히 열어 선택 영역별 대조 분석을 수행한다."
    )
    add_bullets(
        doc,
        [
            "DXF 뷰어: 도면 표시, 영역 선택, 레이어 필터, 숨겨진 항목 삭제",
            "SDF 뷰어: 아이소매트릭 네트워크 표시, 영역 선택, 객체 타입 필터",
            "분석 모달: 헤드 수, 가지/교차 배관 길이, 부속류, 밸브 수 비교",
            "향후 PIPE.001 정밀화 시 explicit mismatch 목록 기반 대조 로직과 연결 가능",
        ],
    )

    doc.add_heading("12. 서버 실행 및 운영 방법", level=1)
    add_numbers(
        doc,
        [
            "프로젝트 폴더에서 python run_server.py 실행",
            "브라우저에서 http://127.0.0.1:5050 접속",
            "외부 PC 접속 시 http://<서버PC_IP>:5050 사용",
            "검증 실행 후 결과 패널과 데이터 테이블 확인",
            "필요 시 엑셀 다운로드 또는 업데이트 기록 확인",
        ],
    )
    doc.add_paragraph("운영 시 주의 사항:")
    add_bullets(
        doc,
        [
            "서버 수정 후에는 기존 run_server.py 프로세스가 남아 있지 않은지 확인 필요",
            "브라우저가 예전 JS/CSS를 들고 있으면 Ctrl+F5 필요",
            "업로드 파일은 data/uploads 아래에 저장됨",
        ],
    )

    doc.add_heading("13. 장애 대응 포인트", level=1)
    add_bullets(
        doc,
        [
            "배관 역할/하류 헤드 수가 비어 보이면 /api/validate 응답과 실행 중 서버 프로세스 버전 확인",
            "REVIEW 때문에 화면이 이상해 보이면 rules.pipe[]와 tables.pipes[*].pipe_rule_results 확인",
            "그래프/텍스트 반영이 안 되면 서버 재시작 후 브라우저 강력 새로고침",
            "CAD 모듈 파싱 오류 시 DXF 파일 형식과 서버 파싱 경로 확인",
        ],
    )

    doc.add_heading("14. 주요 소스 파일 설명", level=1)
    add_bullets(
        doc,
        [
            "대조 서버.py: Flask 엔드포인트, 시각화, SDF 그래프, 엑셀 export",
            "pipenet_validator.py: 전체 파서, 검증 엔진, structured rule 생성",
            "cad_engine.py: DXF 직렬화, 네트워크 엔티티 추출, CAD mismatch helper",
            "head_detector.py: CAD 헤드 탐지 보조",
            "templates/index.html: 메인 UI 골격",
            "static/app.js: 동작/렌더링/모달 제어",
            "static/styles.css: 전체 스타일",
            "data/update_history.json: 업데이트 기록 데이터",
        ],
    )

    doc.add_heading("15. 향후 추가 입력이 필요한 항목", level=1)
    add_bullets(
        doc,
        [
            "CAD 파일: PIPE.001 도면-스케매틱 일치 규칙 고도화",
            "project_meta.json: PIPE.004, PIPE.005 확정 판정",
            "design_policy.json: PIPE.006 회사 정책 기반 판정",
        ],
    )

    if image_path and image_path.exists():
        doc.add_section(WD_SECTION.CONTINUOUS)
        doc.add_heading("부록 A. 참고 화면 이미지", level=1)
        doc.add_paragraph("아래 이미지는 설명서에 참고용으로 삽입한 화면 자료이다.")
        doc.add_picture(str(image_path), width=Inches(6.5))

    doc.add_heading("부록 B. 현재 버전 요약", level=1)
    add_bullets(
        doc,
        [
            "공통 HW 검산 구현 완료",
            "배관 topology 유속 판정 구현 완료",
            "배관 대제목 규칙 엔진(PIPE.001~PIPE.006) 구현 완료",
            "배관 규칙 상태 보기 버튼/모달 구현 완료",
            "REVIEW는 노란색 행 강조 대신 상태값으로만 표시",
        ],
    )

    doc.save(docx_path)


def main() -> None:
    image_candidates = [
        Path(r"C:\Users\admin\Desktop\App\HyperSnap\377.png"),
        Path(r"C:\Users\admin\AppData\Local\Temp\ai-chat-attachment-12243213516581316316.png"),
    ]
    image_path = next((p for p in image_candidates if p.exists()), None)
    out_path = DESKTOP / "PIPENET_검증서버_상세설명서_2026-04-10.docx"
    build_manual(out_path, image_path)
    print(str(out_path))


if __name__ == "__main__":
    main()
