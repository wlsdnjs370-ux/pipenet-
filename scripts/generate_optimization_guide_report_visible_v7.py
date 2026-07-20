from __future__ import annotations

import os
import shutil
import subprocess
from pathlib import Path

import win32com.client
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from generate_optimization_guide_report_v4 import (
    make_conflict_flow,
    make_economy_flow,
    make_engineering_flow,
    make_overall_flow,
)


DESKTOP = Path(os.environ.get("USERPROFILE", r"C:\Users\admin")) / "Desktop"
DOCX = DESKTOP / "PIPENET_2_optimization_guide_visible_v7.docx"
HWP = DESKTOP / "PIPENET_2_optimization_guide_visible_v7.hwp"
HWP_KR = DESKTOP / "PIPENET 수리계산 검증 프로그램 중 2. 설계 최적화 가이드 로직_표그림_가시본_v7.hwp"
HWPX = DESKTOP / "PIPENET_2_optimization_guide_visible_v7.hwpx"

FONT_EA = "함초롬바탕"
FONT_LATIN = "HCR Batang"


def run_font(run, size: int = 10, bold: bool = False) -> None:
    run.font.name = FONT_LATIN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EA)
    run.font.size = Pt(size)
    run.bold = bold


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, text: str, bold: bool = False, fill: str | None = None) -> None:
    cell.text = ""
    if fill:
        shade(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) < 20 else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    run_font(r, 10, bold)
    if fill == "1F2937":
        r.font.color.rgb = RGBColor(255, 255, 255)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell_text(t.rows[0].cells[i], h, True, "1F2937")
        if widths:
            t.rows[0].cells[i].width = Cm(widths[i])
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cell_text(cells[i], val)
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph()


def heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading("", level=level)
    r = p.add_run(text)
    run_font(r, 13 if level == 1 else 11, True)


def para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    run_font(r, 10)


def setup_doc(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EA)
    normal.font.size = Pt(10)
    sec = doc.sections[0]
    sec.top_margin = Cm(1.6)
    sec.bottom_margin = Cm(1.6)
    sec.left_margin = Cm(1.6)
    sec.right_margin = Cm(1.6)


def build_docx() -> None:
    diagrams = [make_overall_flow(), make_engineering_flow(), make_economy_flow(), make_conflict_flow()]
    doc = Document()
    setup_doc(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Analysis Report")
    run_font(r, 20, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PIPENET 수리계산 검진 프로그램 보고서")
    run_font(r, 16, True)

    add_table(
        doc,
        ["구분", "내용"],
        [
            ["보고서 주제", "2. 설계 최적화 가이드 로직"],
            ["문서 목적", "공학적 마찰손실 최적화와 시공사 경제성 확보 방안의 서버 판정 로직 설명"],
            ["작성 방식", "원본 양식의 목차/No/항목명 구조를 기준으로 표와 삽도 중심 작성"],
        ],
        [3.5, 12.0],
    )

    heading(doc, "목차.", 1)
    add_table(
        doc,
        ["No", "항목명"],
        [
            ["1", "문서 목적 및 기본 전제"],
            ["2", "입력 데이터와 파서 구조"],
            ["3", "공통 계산 지표와 수식"],
            ["4", "공학적 마찰손실 최적화 로직"],
            ["5", "Friction Loss Spike Map"],
            ["6", "시공사 경제성 확보 로직"],
            ["7", "Economy Optimization Candidate Map"],
            ["8", "결과 데이터 테이블 표시 구조"],
            ["9", "공학-경제 충돌 및 절충 판단"],
            ["10", "관경 축소 시뮬레이션 흐름"],
            ["11", "한계 및 추가 입력"],
        ],
        [2.0, 13.0],
    )

    heading(doc, "1. 문서 목적 및 기본 전제", 1)
    para(doc, "설계 최적화는 법적 기준을 무시하고 관경을 줄이는 작업이 아니다. 공학안과 경제안 모두 필수 검증 기준을 통과한 설계를 전제로 하며, 이후 안정성 또는 비용 절감 관점의 개선 후보를 분리해 제시한다.")
    add_table(
        doc,
        ["No", "필수 전제", "판정 의미"],
        [
            ["1", "헤드 최소 방수압 0.1 MPa 이상", "말단 방수압 기준 만족 후 최적화 검토 가능"],
            ["2", "헤드 최소 방수량 기준 만족", "노즐별 요구 유량 이상 확보"],
            ["3", "가지배관 유속 6 m/s 이하", "Topology 기준 branch 배관에 적용"],
            ["4", "그 밖의 배관 유속 10 m/s 이하", "교차배관/주배관/50A 초과 배관에 적용"],
            ["5", "재질별 C-Factor 기준 만족", "KSD 계열 C=120, CPVC 계열 C=150"],
            ["6", "특수설비/밸브 등가길이 검증", "FX, AV/PV, PRV 등 입력값 검토"],
            ["7", "Hazen-Williams 재계산 검증", "결과서 Frict. Loss를 수식으로 재현"],
        ],
        [1.2, 6.0, 8.2],
    )
    doc.add_picture(str(diagrams[0]), width=Inches(6.25))

    heading(doc, "2. 입력 데이터와 파서 구조", 1)
    add_table(
        doc,
        ["입력", "주요 추출 필드", "역할"],
        [
            ["결과서 DOCX/PDF", "PIPE CONFIGURATION, FLOW IN PIPES, NOZZLE CONFIGURATION", "압력, 유량, 유속, 마찰손실, 구경, C-Factor의 권위 있는 출력값"],
            ["SDF 파일", "Node, Pipe input/output, Nozzle input, waypoint", "배관망 연결 관계, 하류 헤드 수, 네트워크 맵 표시"],
            ["project_meta.json", "unit_internal_pipe_labels, unit_inlet_pipe_labels", "결과서만으로 알 수 없는 공간 의미 정보 보완"],
            ["design_policy.json", "head_count_min_nominal_by_head_type", "프로젝트별 설계 정책 적용"],
        ],
        [3.0, 5.8, 6.8],
    )

    heading(doc, "3. 공통 계산 지표와 수식", 1)
    add_table(
        doc,
        ["지표", "수식", "해석"],
        [
            ["m당 마찰손실", "friction_loss / length", "길이 영향을 제거한 손실 집중도"],
            ["마찰손실 변화율", "(현재 m당 손실 - 직전 m당 손실) / 직전 m당 손실", "직전 배관 대비 급증 여부"],
            ["유속 여유율", "1 - actual_velocity / velocity_limit", "유속 기준까지 남은 여유"],
            ["압력 여유", "actual_pressure - required_pressure", "압력 안정성 또는 축소 여지"],
            ["등가길이 비율", "(fitting_eq + special_eq) / total_length", "피팅/특수설비 집중도"],
            ["HW 재계산", "6.174e4 * Q^1.85 * L / (C^1.85 * D^4.87) / 0.1", "결과서 손실 재현 검산"],
        ],
        [3.0, 6.2, 6.2],
    )

    heading(doc, "4. 공학적 마찰손실 최적화 로직", 1)
    para(doc, "공학적 최적화는 안정성, 압력 여유, 마찰손실 분산, 유속 안정성을 우선한다. 비용 절감보다 특정 배관에 손실이 집중되는 현상과 말단 압력 불안정을 줄이는 것이 목적이다.")
    add_table(
        doc,
        ["후보 사유", "판정 조건", "권장 조치"],
        [
            ["마찰손실 절대값 과다", "m당 마찰손실이 공학 기준 초과", "관경 상향, 피팅 축소, 경로 단순화"],
            ["마찰손실 변화율 급증", "직전 배관 대비 변화율이 큼", "이전/현재 구간의 구경, 피팅, 특수설비 비교"],
            ["유속 여유 부족", "유속 여유율이 낮아 기준에 근접", "관경 상향 또는 분기/경로 재구성"],
            ["피팅/특수설비 집중", "등가길이 비율이 과다", "엘보/Tee 축소, 특수설비 위치 조정"],
            ["압력 여유 부족", "말단 압력이 요구 기준에 근접", "손실 저감 또는 관경 상향"],
        ],
        [3.8, 5.4, 6.0],
    )
    doc.add_picture(str(diagrams[1]), width=Inches(6.25))

    heading(doc, "5. Friction Loss Spike Map", 1)
    add_table(
        doc,
        ["구분", "적용 기준", "인터랙션"],
        [
            ["표시 목적", "마찰손실 급증 전용 조건만 네트워크에 표시", "빨간 배관 드래그/호버 시 Pipe 번호, 손실, m당 손실, 변화율 표시"],
            ["전용 조건", "m당 마찰손실 > 1.0 kg/cm²/m", "결과 데이터 테이블의 전체 공학 후보와 구분"],
            ["해석", "짧은 배관인데 손실이 크면 피팅/특수설비 집중 가능성", "원인과 조치 방향 동시 표시"],
        ],
        [3.0, 5.4, 6.4],
    )

    heading(doc, "6. 시공사 경제성 확보 로직", 1)
    para(doc, "경제성 최적화는 가능한 싸게 만드는 것이 아니라, 법적/기술 기준을 유지하는 범위에서 과설계를 줄이는 방향이다. 관경 축소 후에도 유속, 말단압력, 헤드 유량, HW 검산 조건을 재확인해야 한다.")
    add_table(
        doc,
        ["경제성 후보", "판정 조건", "권장 조치"],
        [
            ["저유속 과설계", "velocity < 2.0 m/s and nominal_bore > 25A", "한 단계 관경 축소 시뮬레이션"],
            ["압력 여유 과다", "말단압력이 최소 기준보다 과도하게 높음", "0.11~0.12 MPa 수준 목표 검토"],
            ["대구경 밸브 비용", "valve_connected_bore > 100A", "100A 이하 조정 가능성 검토"],
            ["CPVC 대구경", "CPVC nominal_bore > 65A", "50A/65A 복수 라인 분산 검토"],
            ["하류 헤드 수 대비 과대 관경", "downstream_nozzle_count 대비 nominal_bore가 큼", "40A→32A, 32A→25A 등 검토"],
        ],
        [3.8, 5.8, 5.8],
    )
    doc.add_picture(str(diagrams[2]), width=Inches(6.25))

    heading(doc, "7. Economy Optimization Candidate Map", 1)
    add_table(
        doc,
        ["구분", "표시 대상", "주의사항"],
        [
            ["목적", "경제성 후보 배관을 네트워크상에서 확인", "빨간 Spike Map과 목적이 다르므로 혼동 금지"],
            ["대상", "저유속, 압력여유 과다, 대구경 밸브, CPVC 대구경", "최종 축소 여부는 재계산 통과 후 결정"],
            ["조작", "Friction Loss Spike Map과 동일하게 줌/드래그 가능", "시각화는 후보 탐색 도구이며 자동 설계 변경은 아님"],
        ],
        [3.0, 5.8, 6.6],
    )

    heading(doc, "8. 결과 데이터 테이블 표시 구조", 1)
    add_table(
        doc,
        ["열 그룹", "주요 필드", "클릭 시 설명"],
        [
            ["기본 배관 정보", "Pipe, 구경, 실내경, 길이, 유량, 유속", "원자료 출처와 단위 설명"],
            ["HW 검산 상세", "C-Factor, 총등가길이, 결과서 손실, 재계산 손실, 오차", "수식, 허용오차, PASS/FAIL 근거"],
            ["Topology 유속", "배관 역할, 하류 헤드 수, 하류 교차분기, 유속 기준", "branch/other 판정 이유"],
            ["공학 후보", "마찰손실, 변화율, 유속여유, 피팅집중, 압력여유", "각 사유별 계산값과 조치"],
            ["경제 후보", "저유속, 압력여유 과다, 대구경 밸브, CPVC 대구경", "축소 시뮬레이션 필요 조건"],
        ],
        [3.0, 6.0, 6.6],
    )

    heading(doc, "9. 공학-경제 충돌 및 절충 판단", 1)
    add_table(
        doc,
        ["충돌 유형", "공학 해석", "경제 해석", "절충안"],
        [
            ["손실 큰 배관", "구경 상향 또는 피팅 축소 필요", "구경 상향은 비용 증가", "관경 변경 전 피팅/경로부터 점검"],
            ["저유속 대구경", "압력 안정성에는 유리", "과설계 가능성", "축소 시뮬레이션 후 기준 재확인"],
            ["CPVC 대구경", "C=150으로 손실 저감 가능", "대구경 CPVC 비용 상승", "복수 소구경 라인 또는 재질 대안 비교"],
            ["대구경 밸브", "손실과 유지관리 측면에서 안정", "밸브 단가 상승", "100A 이하 조정 가능성 검토"],
        ],
        [3.0, 4.2, 4.0, 4.8],
    )
    doc.add_picture(str(diagrams[3]), width=Inches(6.25))

    heading(doc, "10. 관경 축소 시뮬레이션 흐름", 1)
    add_table(
        doc,
        ["단계", "처리 내용", "통과 조건"],
        [
            ["1", "현재 설계 압력 여유와 유속 여유 확인", "필수 기준 PASS 상태"],
            ["2", "관경 한 단계 축소 가정", "125A→100A, 100A→80A, 80A→65A 등"],
            ["3", "Hazen-Williams로 마찰손실 재계산", "총등가길이와 C-Factor 반영"],
            ["4", "유속 기준 재검토", "가지 ≤ 6 m/s, 그 밖 ≤ 10 m/s"],
            ["5", "말단압력 및 헤드 유량 재검토", "방수압 ≥ 0.1 MPa, 유량 기준 만족"],
            ["6", "통과 시 경제성 후보로 제시", "자동 확정이 아니라 설계자 검토 대상으로 출력"],
        ],
        [1.4, 7.8, 6.2],
    )

    heading(doc, "11. 한계 및 추가 입력", 1)
    add_table(
        doc,
        ["항목", "현재 한계", "보완 입력"],
        [
            ["세대 내부 CPVC", "결과서만으로 공간 구분 불가", "project_meta.json 또는 CAD zone 분류"],
            ["세대 유입 65A", "어떤 배관이 세대 유입인지 결과서만으로 확정 불가", "unit_inlet_pipe_labels"],
            ["헤드 수별 관경 정책", "법정 기준이 아니라 회사/프로젝트 정책", "design_policy.json"],
            ["도면-스케매틱 일치", "CAD 레이어/축척/좌표계 정리가 필요", "정제 DXF, transform metadata"],
            ["경제성 금액 산출", "현재는 후보 로직 중심", "자재 단가표, 밸브 단가표, 시공 단가표"],
        ],
        [3.2, 6.2, 6.0],
    )

    doc.save(DOCX)


def convert_to_hwp() -> None:
    hwp = win32com.client.DispatchEx("HWPFrame.HwpObject")
    try:
        try:
            hwp.RegisterModule("FilePathCheckDLL", "FilePathCheckerModule")
        except Exception:
            pass
        if not hwp.Open(str(DOCX)):
            raise RuntimeError("DOCX open failed")
        hwp.HAction.GetDefault("FileSaveAs_S", hwp.HParameterSet.HFileOpenSave.HSet)
        hwp.HParameterSet.HFileOpenSave.filename = str(HWP)
        hwp.HParameterSet.HFileOpenSave.Format = "HWP"
        hwp.HParameterSet.HFileOpenSave.Attributes = 0
        if not hwp.HAction.Execute("FileSaveAs_S", hwp.HParameterSet.HFileOpenSave.HSet):
            raise RuntimeError("HWP save failed")
        hwp.HAction.GetDefault("FileSaveAs_S", hwp.HParameterSet.HFileOpenSave.HSet)
        hwp.HParameterSet.HFileOpenSave.filename = str(HWPX)
        hwp.HParameterSet.HFileOpenSave.Format = "HWPX"
        hwp.HParameterSet.HFileOpenSave.Attributes = 0
        hwp.HAction.Execute("FileSaveAs_S", hwp.HParameterSet.HFileOpenSave.HSet)
    finally:
        hwp.Quit()
    shutil.copy2(HWP, HWP_KR)


def validate() -> None:
    for p in [DOCX, HWP, HWP_KR, HWPX]:
        print(f"{p.name}: exists={p.exists()} size={p.stat().st_size if p.exists() else 0}")
    text = subprocess.check_output(["hwp5txt.exe", str(HWP)], text=True, encoding="utf-8", errors="replace", timeout=30)
    markers = ["Analysis Report", "목차", "공학적 마찰손실", "시공사 경제성", "Friction Loss Spike Map"]
    print("MARKERS=" + ", ".join(f"{m}:{m in text}" for m in markers))
    print("QUESTION_MARK_COUNT=", text.count("?"))
    hwp = win32com.client.DispatchEx("HWPFrame.HwpObject")
    try:
        try:
            hwp.RegisterModule("FilePathCheckDLL", "FilePathCheckerModule")
        except Exception:
            pass
        print("COM_OPEN=", bool(hwp.Open(str(HWP))))
    finally:
        hwp.Quit()


if __name__ == "__main__":
    build_docx()
    convert_to_hwp()
    validate()
