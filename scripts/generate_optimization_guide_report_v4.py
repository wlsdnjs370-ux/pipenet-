from __future__ import annotations

import os
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DESKTOP = Path(os.environ.get("USERPROFILE", r"C:\Users\admin")) / "Desktop"
OUT_DIR = ROOT / "docs" / "generated_reports"
OUT_DIR.mkdir(parents=True, exist_ok=True)

DOCX_PATH = DESKTOP / "PIPENET_2_optimization_guide_table_flowchart_v4.docx"
HWP_PATH = DESKTOP / "PIPENET_2_optimization_guide_table_flowchart_v4.hwp"
KOREAN_HWP_PATH = DESKTOP / "PIPENET 수리계산 검증 프로그램 중 2. 설계 최적화 가이드 로직_표_플로우차트_완성본.hwp"
TEMPLATE_PATH = DESKTOP / "PIPENET 수리계산 검증 프로그램 중 2. 설계 최적화 가이드 로직.hwp"


def _font_path() -> str:
    candidates = [
        r"C:\Windows\Fonts\malgun.ttf",
        r"C:\Windows\Fonts\malgunbd.ttf",
        r"C:\Windows\Fonts\arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return candidate
    return ""


FONT_PATH = _font_path()


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    path = r"C:\Windows\Fonts\malgunbd.ttf" if bold and Path(r"C:\Windows\Fonts\malgunbd.ttf").exists() else FONT_PATH
    if path:
        return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_width: int) -> list[str]:
    lines: list[str] = []
    for raw_line in text.split("\n"):
        words = raw_line.split(" ")
        line = ""
        for word in words:
            candidate = word if not line else f"{line} {word}"
            if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width:
                line = candidate
            else:
                if line:
                    lines.append(line)
                line = word
        if line:
            lines.append(line)
    return lines or [""]


def draw_box(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], title: str, body: str, fill: str, outline: str) -> None:
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=16, fill=fill, outline=outline, width=3)
    title_font = load_font(24, bold=True)
    body_font = load_font(18)
    draw.text((x1 + 18, y1 + 14), title, font=title_font, fill="#101820")
    y = y1 + 52
    for line in wrap_text(draw, body, body_font, x2 - x1 - 36):
        draw.text((x1 + 18, y), line, font=body_font, fill="#263238")
        y += 25


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = "#4b5563") -> None:
    draw.line([start, end], fill=color, width=4)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        points = [(ex, ey), (ex - direction * 18, ey - 10), (ex - direction * 18, ey + 10)]
    else:
        direction = 1 if ey > sy else -1
        points = [(ex, ey), (ex - 10, ey - direction * 18), (ex + 10, ey - direction * 18)]
    draw.polygon(points, fill=color)


def make_overall_flow() -> Path:
    img = Image.new("RGB", (1800, 980), "white")
    draw = ImageDraw.Draw(img)
    draw.text((60, 36), "설계 최적화 가이드 전체 데이터 흐름", font=load_font(34, True), fill="#111827")
    boxes = [
        ((70, 130, 410, 300), "입력", "결과서 DOCX/PDF\nSDF 배관망\n선택 입력: project_meta, design_policy", "#f8fafc", "#111827"),
        ((520, 130, 860, 300), "파서", "배관, 헤드, 유속, 압력, C-Factor,\nHW 검산, 토폴로지 데이터 구조화", "#eef2ff", "#1d4ed8"),
        ((970, 130, 1310, 300), "필수 기준", "법적/기술 기준 PASS 여부 확인\nPASS 전제에서만 최적화 제안", "#ecfdf5", "#047857"),
        ((1420, 130, 1760, 300), "리포트", "공학 최적화\n경제성 최적화\n충돌/절충 후보", "#fff7ed", "#c2410c"),
        ((295, 475, 735, 690), "공학 최적화", "m당 마찰손실, 변화율,\n유속 여유, 압력 여유,\n피팅/특수설비 집중도 평가", "#eff6ff", "#2563eb"),
        ((1065, 475, 1505, 690), "경제성 최적화", "저유속 과설계, 대구경 밸브,\nCPVC 대구경, 관경 축소 시뮬레이션,\n하류 헤드 수 대비 과대 관경 평가", "#f0fdf4", "#16a34a"),
        ((685, 795, 1115, 925), "출력", "결과 데이터 테이블, Friction Loss Spike Map,\nEconomy Optimization Candidate Map,\n구간별 조치 카드와 근거", "#fafafa", "#525252"),
    ]
    for spec in boxes:
        draw_box(draw, *spec)
    arrow(draw, (410, 215), (520, 215))
    arrow(draw, (860, 215), (970, 215))
    arrow(draw, (1310, 215), (1420, 215))
    arrow(draw, (1140, 300), (515, 475))
    arrow(draw, (1140, 300), (1285, 475))
    arrow(draw, (515, 690), (815, 795))
    arrow(draw, (1285, 690), (985, 795))
    path = OUT_DIR / "optimization_overall_flow.png"
    img.save(path)
    return path


def make_engineering_flow() -> Path:
    img = Image.new("RGB", (1600, 880), "white")
    draw = ImageDraw.Draw(img)
    draw.text((60, 36), "공학적 마찰손실 최적화 판정 흐름", font=load_font(34, True), fill="#111827")
    boxes = [
        ((70, 130, 390, 310), "배관별 지표", "friction_loss / length\n손실 변화율\n유속 여유율\n압력 여유\n등가길이 비율", "#eff6ff", "#2563eb"),
        ((520, 130, 840, 310), "급증 판정", "Friction Loss Spike Map은\nm당 마찰손실 > 1.0 kg/cm²/m\n전용 조건만 빨강 표시", "#fee2e2", "#dc2626"),
        ((970, 130, 1290, 310), "후보 세분화", "마찰손실, 변화율,\n유속여유, 피팅집중,\n압력여유 사유를 분리", "#fefce8", "#ca8a04"),
        ((640, 480, 960, 670), "조치 제안", "구경 상향\n피팅 축소\n경로 단순화\n특수설비 위치 조정\n루프/분산 배관 검토", "#f8fafc", "#111827"),
    ]
    for spec in boxes:
        draw_box(draw, *spec)
    arrow(draw, (390, 220), (520, 220))
    arrow(draw, (840, 220), (970, 220))
    arrow(draw, (1130, 310), (850, 480))
    arrow(draw, (680, 310), (750, 480))
    path = OUT_DIR / "optimization_engineering_flow.png"
    img.save(path)
    return path


def make_economy_flow() -> Path:
    img = Image.new("RGB", (1600, 880), "white")
    draw = ImageDraw.Draw(img)
    draw.text((60, 36), "시공사 경제성 확보 판정 흐름", font=load_font(34, True), fill="#111827")
    boxes = [
        ((70, 130, 390, 310), "경제성 후보", "저유속 + 25A 초과\n압력 여유 과다\n대구경 밸브\nCPVC 65A 초과\n하류 헤드 수 대비 과대 관경", "#f0fdf4", "#16a34a"),
        ((520, 130, 840, 310), "축소 시뮬레이션", "관경 한 단계 축소 가정\nHW 손실 재계산\n유속 기준 재확인\n말단압력 재확인", "#ecfdf5", "#047857"),
        ((970, 130, 1290, 310), "통과 조건", "가지 ≤ 6 m/s\n그 밖 ≤ 10 m/s\n헤드 압력 ≥ 0.1 MPa\n헤드 유량 기준 만족", "#f8fafc", "#111827"),
        ((640, 480, 960, 670), "경제 제안", "관경 축소 검토\n밸브 100A 이하 검토\nCPVC 대구경 회피\n복수 소구경 라인 분산", "#fff7ed", "#c2410c"),
    ]
    for spec in boxes:
        draw_box(draw, *spec)
    arrow(draw, (390, 220), (520, 220))
    arrow(draw, (840, 220), (970, 220))
    arrow(draw, (1130, 310), (850, 480))
    arrow(draw, (680, 310), (750, 480))
    path = OUT_DIR / "optimization_economy_flow.png"
    img.save(path)
    return path


def make_conflict_flow() -> Path:
    img = Image.new("RGB", (1600, 760), "white")
    draw = ImageDraw.Draw(img)
    draw.text((60, 36), "공학 관점과 경제 관점의 충돌/절충 판단", font=load_font(34, True), fill="#111827")
    boxes = [
        ((90, 150, 440, 350), "공학 관점", "안정성, 압력 여유,\n마찰손실 분산,\n유속 안정성 우선\n비용 증가 가능", "#eff6ff", "#2563eb"),
        ((1160, 150, 1510, 350), "경제 관점", "기준 만족 범위에서\n관경, 밸브, 재질,\n시공비 절감\n여유 감소 가능", "#f0fdf4", "#16a34a"),
        ((610, 150, 990, 350), "충돌 구간", "공학: 구경 상향 후보\n경제: 구경 축소 후보\n또는 압력여유/유속여유 판단 상충", "#fef2f2", "#dc2626"),
        ((610, 500, 990, 650), "절충안", "관경 변경 전 피팅/경로 점검\n축소 시뮬레이션 후 기준 재확인\n위험구간은 공학 우선", "#fafafa", "#111827"),
    ]
    for spec in boxes:
        draw_box(draw, *spec)
    arrow(draw, (440, 250), (610, 250))
    arrow(draw, (1160, 250), (990, 250))
    arrow(draw, (800, 350), (800, 500))
    path = OUT_DIR / "optimization_conflict_flow.png"
    img.save(path)
    return path


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if "\n" not in text and len(text) < 24 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(9)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, color="FFFFFF")
        set_cell_shading(table.rows[0].cells[i], "1F2937")
        if widths:
            table.rows[0].cells[i].width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "맑은 고딕"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")


def add_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(10)


def set_doc_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    style.font.size = Pt(10)
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Analysis Report")
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(22)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PIPENET 수리계산 검진 프로그램 보고서")
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(18)
    run.bold = True

    add_table(
        doc,
        ["구분", "내용"],
        [
            ["보고서 주제", "2. 설계 최적화 가이드 로직"],
            ["작성 목적", "공학적 마찰손실 최적화와 시공사 경제성 확보 방안을 서버 판정 로직 기준으로 정리"],
            ["문서 형식", "표, 수식, 플로우차트, 네트워크 맵 설명 중심"],
            ["참조 양식", "원본 HWP의 목차 / No / 항목명 테이블 구조"],
        ],
        [4.0, 12.5],
    )


def add_toc(doc: Document) -> None:
    add_heading(doc, "목차.", 1)
    rows = [
        ["1", "문서 목적 및 기본 전제"],
        ["2", "입력 데이터와 파서 구조"],
        ["3", "공통 계산 지표와 수식"],
        ["4", "공학적 마찰손실 최적화 로직"],
        ["5", "Friction Loss Spike Map"],
        ["6", "시공사 경제성 확보 로직"],
        ["7", "Economy Optimization Candidate Map"],
        ["8", "결과 데이터 테이블 표시 구조"],
        ["9", "공학-경제 충돌 및 절충 판단"],
        ["10", "관경 축소 시뮬레이션 흐름"],
        ["11", "한계 및 추가 입력"],
    ]
    add_table(doc, ["No", "항목명"], rows, [2.0, 14.0])


def build_docx() -> None:
    diagrams = [make_overall_flow(), make_engineering_flow(), make_economy_flow(), make_conflict_flow()]
    doc = Document()
    set_doc_style(doc)
    add_cover(doc)
    add_toc(doc)

    add_heading(doc, "1. 문서 목적 및 기본 전제", 1)
    add_note(doc, "설계 최적화는 법적 기준을 무시하고 관경을 줄이는 작업이 아니다. 공학안과 경제안 모두 필수 검증 기준을 통과한 설계를 전제로 하며, 이후 안정성 또는 비용 절감 관점의 개선 후보를 분리해 제시한다.")
    add_table(
        doc,
        ["No", "필수 전제", "판정 의미"],
        [
            ["1", "헤드 최소 방수압 0.1 MPa 이상", "말단 방수압이 법적 최소 기준을 만족해야 최적화 검토 가능"],
            ["2", "헤드 최소 방수량 기준 만족", "표준헤드 등 노즐별 요구 유량 이상 확보"],
            ["3", "가지배관 유속 6 m/s 이하", "Topology 기준 branch 판정 배관에 적용"],
            ["4", "그 밖의 배관 유속 10 m/s 이하", "교차배관/주배관/50A 초과 배관에 적용"],
            ["5", "재질별 C-Factor 기준 만족", "KSD 계열 C=120, CPVC 계열 C=150"],
            ["6", "특수설비/밸브/감압밸브 검증", "등가길이, 설정압력, 장치 입력값이 설계 기준과 충돌하지 않는지 확인"],
            ["7", "Hazen-Williams 마찰손실 재계산", "결과서 Frict. Loss를 수식으로 재현해 계산 신뢰성 확인"],
        ],
        [1.2, 6.2, 9.0],
    )
    doc.add_picture(str(diagrams[0]), width=Inches(6.7))

    add_heading(doc, "2. 입력 데이터와 파서 구조", 1)
    add_table(
        doc,
        ["입력", "사용 위치", "주요 추출 필드", "역할"],
        [
            ["결과서 DOCX/PDF", "기본 검증 및 최적화", "PIPE CONFIGURATION, DESIGNED DIAMETERS & FLOWRATES, FLOW IN PIPES, NOZZLE CONFIGURATION", "압력, 유량, 유속, 마찰손실, 구경, C-Factor의 권위 있는 출력값"],
            ["SDF 파일", "토폴로지 및 맵 시각화", "Node, Pipe input/output, Nozzle input, waypoint, 특수설비", "배관망 연결 관계, 하류 헤드 수, 네트워크 맵 표시"],
            ["project_meta.json", "세대 내부/세대 유입 판정", "unit_internal_pipe_labels, unit_inlet_pipe_labels", "결과서만으로 알 수 없는 공간 의미 정보 보완"],
            ["design_policy.json", "헤드 수별 최소 관경 정책", "head_count_min_nominal_by_head_type", "회사/프로젝트별 설계 정책을 하드코딩하지 않고 적용"],
        ],
        [3.0, 3.5, 5.6, 4.5],
    )

    add_heading(doc, "3. 공통 계산 지표와 수식", 1)
    add_table(
        doc,
        ["지표", "수식", "데이터 소스", "해석"],
        [
            ["m당 마찰손실", "friction_loss / length", "FLOW IN PIPES Frict. Loss + PIPE CONFIGURATION Length", "배관 길이 영향을 제거해 손실 집중도를 비교"],
            ["마찰손실 변화율", "(현재 m당 손실 - 직전 m당 손실) / 직전 m당 손실", "배관 순번별 m당 마찰손실", "직전 배관 대비 손실 급증 여부 판단"],
            ["유속 여유율", "1 - actual_velocity / velocity_limit", "FLOW IN PIPES Velocity + topology role", "유속 기준까지 남은 여유 확인"],
            ["압력 여유", "actual_pressure - required_pressure", "NOZZLE/FLOW 압력값", "말단 압력 안정성 또는 관경 축소 가능성 판단"],
            ["등가길이 비율", "(fitting_eq + special_eq) / total_length", "PIPE CONFIGURATION + SPECIAL EQUIPMENT", "피팅/특수설비가 손실을 집중시키는지 확인"],
            ["HW 재계산 손실", "6.174e4 * Q^1.85 * L / (C^1.85 * D^4.87) / 0.1", "Q=LPM, L=총등가길이, D=실내경(mm)", "결과서 Frict. Loss 재현 검산"],
        ],
        [3.2, 5.4, 4.2, 3.8],
    )

    add_heading(doc, "4. 공학적 마찰손실 최적화 로직", 1)
    add_note(doc, "공학적 최적화는 안정성, 압력 여유, 마찰손실 분산, 유속 안정성을 우선한다. 비용 절감보다 특정 배관에 손실이 집중되는 현상과 말단 압력 불안정을 줄이는 것이 목적이다.")
    add_table(
        doc,
        ["후보 사유", "판정 조건", "출력 색상/표시", "권장 조치"],
        [
            ["마찰손실 절대값 과다", "m당 마찰손실이 공학 기준을 초과", "공학 최적화 후보", "관경 상향, 피팅 축소, 경로 단순화"],
            ["마찰손실 변화율 급증", "직전 배관 대비 m당 손실 변화율이 큰 구간", "후보 사유: 변화율", "이전/현재 구간의 구경, 피팅, 특수설비 비교"],
            ["유속 여유 부족", "유속 여유율이 낮아 기준에 근접", "후보 사유: 유속여유", "관경 상향 또는 분기/경로 재구성"],
            ["피팅/특수설비 집중", "등가길이 비율이 과다", "후보 사유: 피팅집중", "엘보/Tee 축소, 특수설비 위치 조정"],
            ["압력 여유 부족", "말단 압력이 요구 기준에 근접", "후보 사유: 압력여유", "손실 저감 또는 관경 상향"],
        ],
        [3.6, 5.0, 3.2, 4.6],
    )
    doc.add_picture(str(diagrams[1]), width=Inches(6.7))

    add_heading(doc, "5. Friction Loss Spike Map", 1)
    add_table(
        doc,
        ["구분", "적용 기준", "표시 방식", "인터랙션"],
        [
            ["표시 목적", "마찰손실 급증 전용 조건만 네트워크에 표시", "빨간 배관", "마우스 드래그/호버 시 Pipe 번호, 마찰손실, m당 손실, 변화율 표시"],
            ["전용 조건", "m당 마찰손실 > 1.0 kg/cm²/m", "Spike Map에만 적용", "결과 데이터 테이블의 전체 공학 후보와 구분"],
            ["해석", "짧은 배관인데 손실이 크면 피팅/특수설비 집중 가능성, 긴 배관이면 길이 영향 가능성", "상세 카드", "원인과 조치 방향 동시 표시"],
        ],
        [3.0, 5.3, 3.3, 5.0],
    )

    add_heading(doc, "6. 시공사 경제성 확보 로직", 1)
    add_note(doc, "경제성 최적화는 가능한 싸게 만드는 것이 아니라, 법적/기술 기준을 유지하는 범위에서 과설계를 줄이는 방향이다. 관경 축소 후에도 유속, 말단압력, 헤드 유량, HW 검산 조건을 재확인해야 한다.")
    add_table(
        doc,
        ["경제성 후보", "판정 조건", "검토 의미", "권장 조치"],
        [
            ["저유속 과설계", "velocity < 2.0 m/s and nominal_bore > 25A", "현재 관경 대비 유량이 작아 자재비 절감 여지", "한 단계 관경 축소 시뮬레이션"],
            ["압력 여유 과다", "말단압력 또는 헤드압력이 최소 기준보다 과도하게 높음", "압력 여유를 관경 축소로 일부 회수 가능", "목표 압력 0.11~0.12 MPa 수준 검토"],
            ["대구경 밸브 비용", "valve_connected_bore > 100A", "밸브 단가 급증 구간", "수리계산 여유 시 100A 이하 조정"],
            ["CPVC 대구경", "CPVC nominal_bore > 65A", "CPVC 대구경 자재비 상승 가능성", "50A/65A 복수 라인 분산 검토"],
            ["하류 헤드 수 대비 과대 관경", "downstream_nozzle_count 대비 nominal_bore가 큼", "물량이 많은 가지/말단 구간 최적화", "40A→32A, 32A→25A 등 검토"],
        ],
        [3.3, 4.8, 4.2, 4.3],
    )
    doc.add_picture(str(diagrams[2]), width=Inches(6.7))

    add_heading(doc, "7. Economy Optimization Candidate Map", 1)
    add_table(
        doc,
        ["구분", "표시 대상", "표시 방식", "주의사항"],
        [
            ["목적", "경제성 후보 배관을 네트워크상에서 확인", "녹색 계열 후보 표시", "빨간 Spike Map과 목적이 다르므로 혼동 금지"],
            ["대상", "저유속, 압력여유 과다, 대구경 밸브, CPVC 대구경, 하류 헤드 수 대비 과대 관경", "후보별 카드/툴팁", "최종 축소 여부는 재계산 통과 후 결정"],
            ["조작", "Friction Loss Spike Map과 동일하게 줌/드래그 가능", "+/- 버튼 및 마우스 이동", "시각화는 후보 탐색 도구이며 자동 설계 변경은 아님"],
        ],
        [3.0, 5.4, 3.8, 4.0],
    )

    add_heading(doc, "8. 결과 데이터 테이블 표시 구조", 1)
    add_table(
        doc,
        ["열 그룹", "주요 필드", "의미", "클릭 시 설명"],
        [
            ["기본 배관 정보", "Pipe, 구경, 실내경, 길이, 유량, 유속", "결과서와 SDF에서 파싱한 원자료", "원자료 출처와 단위 설명"],
            ["HW 검산 상세", "C-Factor, 총등가길이, 결과서 손실, 재계산 손실, 오차", "Hazen-Williams 공식 재현 여부", "수식, 허용오차, PASS/FAIL 근거"],
            ["Topology 유속", "배관 역할, 하류 헤드 수, 하류 교차분기, 유속 기준", "가지/그 밖 배관 판정 근거", "branch/other 판정 이유"],
            ["공학 후보", "마찰손실, 변화율, 유속여유, 피팅집중, 압력여유", "안정성 개선 후보 사유 세분화", "각 사유별 계산값과 조치"],
            ["경제 후보", "저유속, 압력여유 과다, 대구경 밸브, CPVC 대구경", "비용 절감 검토 후보", "축소 시뮬레이션 필요 조건"],
        ],
        [3.2, 4.4, 4.5, 4.5],
    )

    add_heading(doc, "9. 공학-경제 충돌 및 절충 판단", 1)
    add_table(
        doc,
        ["충돌 유형", "공학 해석", "경제 해석", "절충안"],
        [
            ["손실이 큰 배관", "관경 상향 또는 피팅 축소 필요", "관경 상향은 비용 증가", "관경 변경 전 피팅/경로부터 점검"],
            ["저유속 대구경 배관", "압력 안정성에는 유리", "과설계 가능성", "축소 시뮬레이션 후 유속/압력 재검증"],
            ["CPVC 대구경", "C=150으로 손실 저감 가능", "대구경 CPVC 비용 상승", "복수 소구경 라인 또는 재질 대안 비교"],
            ["대구경 밸브", "손실과 유지관리 측면에서 안정적일 수 있음", "밸브 단가 상승", "100A 이하 조정 가능성을 압력 여유로 검토"],
        ],
        [3.2, 4.2, 4.2, 4.8],
    )
    doc.add_picture(str(diagrams[3]), width=Inches(6.7))

    add_heading(doc, "10. 관경 축소 시뮬레이션 흐름", 1)
    add_table(
        doc,
        ["단계", "처리 내용", "통과 조건"],
        [
            ["1", "현재 설계 압력 여유와 유속 여유 확인", "필수 기준 PASS 상태"],
            ["2", "관경 한 단계 축소 가정", "예: 125A→100A, 100A→80A, 80A→65A, 40A→32A"],
            ["3", "Hazen-Williams로 마찰손실 재계산", "총등가길이와 C-Factor 반영"],
            ["4", "유속 기준 재검토", "가지 ≤ 6 m/s, 그 밖 ≤ 10 m/s"],
            ["5", "말단압력 및 헤드 유량 재검토", "방수압 ≥ 0.1 MPa, 유량 기준 만족"],
            ["6", "통과 시 경제성 후보로 제시", "자동 확정이 아니라 설계자 검토 대상으로 출력"],
        ],
        [1.5, 8.0, 6.5],
    )

    add_heading(doc, "11. 한계 및 추가 입력", 1)
    add_table(
        doc,
        ["항목", "현재 한계", "보완 입력"],
        [
            ["세대 내부 CPVC", "결과서만으로 공간 구분 불가", "project_meta.json 또는 CAD zone 분류"],
            ["세대 유입 65A", "어떤 배관이 세대 유입인지 결과서만으로 확정 불가", "unit_inlet_pipe_labels"],
            ["헤드 수별 관경 정책", "법정 기준이 아니라 회사/프로젝트 정책", "design_policy.json"],
            ["도면-스케매틱 일치", "CAD 레이어/축척/좌표계 정리가 필요", "정제 DXF, transform metadata"],
            ["경제성 금액 산출", "현재는 후보 로직 중심", "자재 단가표, 밸브 단가표, 시공 단가표"],
        ],
        [3.5, 6.0, 6.5],
    )

    doc.save(DOCX_PATH)


def convert_docx_to_hwp() -> bool:
    try:
        import win32com.client

        hwp = win32com.client.DispatchEx("HWPFrame.HwpObject")
        try:
            hwp.RegisterModule("FilePathCheckDLL", "FilePathCheckerModule")
        except Exception:
            pass
        if not hwp.Open(str(DOCX_PATH)):
            hwp.Quit()
            return False
        hwp.HAction.GetDefault("FileSaveAs_S", hwp.HParameterSet.HFileOpenSave.HSet)
        hwp.HParameterSet.HFileOpenSave.filename = str(HWP_PATH)
        hwp.HParameterSet.HFileOpenSave.Format = "HWP"
        hwp.HParameterSet.HFileOpenSave.Attributes = 0
        ok = bool(hwp.HAction.Execute("FileSaveAs_S", hwp.HParameterSet.HFileOpenSave.HSet))
        hwp.Quit()
        if ok and HWP_PATH.exists():
            shutil.copy2(HWP_PATH, KOREAN_HWP_PATH)
        return ok and HWP_PATH.exists()
    except Exception as exc:
        print(f"HWP conversion failed: {exc}")
        return False


def validate_outputs() -> None:
    print(f"DOCX={DOCX_PATH} exists={DOCX_PATH.exists()} size={DOCX_PATH.stat().st_size if DOCX_PATH.exists() else 0}")
    print(f"HWP={HWP_PATH} exists={HWP_PATH.exists()} size={HWP_PATH.stat().st_size if HWP_PATH.exists() else 0}")
    print(f"KOREAN_HWP={KOREAN_HWP_PATH} exists={KOREAN_HWP_PATH.exists()} size={KOREAN_HWP_PATH.stat().st_size if KOREAN_HWP_PATH.exists() else 0}")
    if HWP_PATH.exists():
        try:
            out = subprocess.check_output(["hwp5txt.exe", str(HWP_PATH)], text=True, encoding="utf-8", errors="replace", timeout=30)
            markers = ["목차", "공학적 마찰손실", "시공사 경제성", "Friction Loss Spike Map"]
            print("TEXT_MARKERS=" + ", ".join(f"{m}:{m in out}" for m in markers))
        except Exception as exc:
            print(f"hwp5txt validation failed: {exc}")


if __name__ == "__main__":
    build_docx()
    converted = convert_docx_to_hwp()
    print(f"converted={converted}")
    validate_outputs()
