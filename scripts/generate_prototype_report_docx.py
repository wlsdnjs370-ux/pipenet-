"""Remote 30 프로토타입 모듈 — 일반인용 .docx 보고서 + 삽도 자동 생성.

생성물:
  data/prototype_report/figures/*.png — matplotlib 삽도
  data/prototype_report/Remote30_프로토타입_원리_보고서.docx — 워드 보고서
"""

from __future__ import annotations

import io
import sys
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.patches as patches
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Inches, Pt, RGBColor

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

PROJECT = Path(r"C:\Users\admin\PycharmProjects\JupyterProject")
OUT_DIR = PROJECT / "data" / "prototype_report"
FIG_DIR = OUT_DIR / "figures"
OUT_DIR.mkdir(parents=True, exist_ok=True)
FIG_DIR.mkdir(parents=True, exist_ok=True)

# Windows 한글 폰트 (Malgun Gothic)
plt.rcParams["font.family"] = "Malgun Gothic"
plt.rcParams["axes.unicode_minus"] = False


# ────────────────────────────────────────────────────────────────────────────
# 삽도 생성
# ────────────────────────────────────────────────────────────────────────────


def fig_pipeline_flow():
    """그림 1 — 6 stage 흐름도."""
    fig, ax = plt.subplots(figsize=(11, 4.5))
    ax.set_xlim(0, 14); ax.set_ylim(0, 5)
    ax.axis("off")
    stages = [
        ("0\nDXF\n파싱", "#6366f1"),
        ("1\n배관망\n추출", "#0ea5e9"),
        ("2\n헤드\n인식", "#f97316"),
        ("3\n배관망\n그래프", "#06b6d4"),
        ("4\n30 헤드\n선정", "#ef4444"),
        ("5\n입력\n테이블", "#f59e0b"),
        ("6\nSDF\nemit", "#10b981"),
    ]
    x = 0.3; w = 1.7; h = 2.2
    for i, (label, color) in enumerate(stages):
        rect = patches.FancyBboxPatch((x, 1.5), w, h, boxstyle="round,pad=0.05",
                                       linewidth=1.6, edgecolor=color, facecolor=color, alpha=0.18)
        ax.add_patch(rect)
        ax.text(x + w/2, 1.5 + h/2, label, ha="center", va="center",
                fontsize=11, fontweight="bold", color=color)
        if i < len(stages) - 1:
            ax.annotate("", xy=(x + w + 0.18, 1.5 + h/2), xytext=(x + w + 0.05, 1.5 + h/2),
                        arrowprops=dict(arrowstyle="->", color="#475569", lw=1.5))
        x += w + 0.22
    # 사용자 편집 단계 표시
    ax.add_patch(patches.FancyBboxPatch((4.3, 0.2), 6.0, 1.0, boxstyle="round,pad=0.05",
                                         linewidth=1.4, edgecolor="#7c3aed",
                                         facecolor="#7c3aed", alpha=0.15, linestyle="--"))
    ax.text(7.3, 0.7, "사용자 편집 (헤드 추가/삭제, 영역 지정, 알람밸브 좌표)",
            ha="center", va="center", fontsize=10, fontweight="bold", color="#5b21b6")
    ax.annotate("", xy=(7.3, 1.45), xytext=(7.3, 1.25),
                arrowprops=dict(arrowstyle="->", color="#5b21b6", lw=1.5, linestyle="--"))
    ax.text(7.0, 4.7, "전체 6-Stage 파이프라인", ha="center", va="center",
            fontsize=13, fontweight="bold")
    plt.tight_layout()
    path = FIG_DIR / "fig1_pipeline_flow.png"
    plt.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close()
    return path


def fig_snap_graph_concept():
    """그림 2 — Snap + bridge 개념."""
    fig, axes = plt.subplots(1, 3, figsize=(12, 4))
    np.random.seed(0)
    # (a) 원본 LINE 끝점들 (정밀도 오차 있음)
    ax = axes[0]
    raw_lines = [
        ((0, 0), (1, 0)), ((1.02, 0.01), (2, 0)),
        ((2.01, -0.01), (3, 0)), ((1, 0.02), (1, 1)),
        ((1.01, 1), (1, 2)), ((3, 0.01), (3, 1)),
        ((5, 0), (6, 0)), ((6, 0), (6, 1)),  # 다른 component
    ]
    for (a, b) in raw_lines:
        ax.plot([a[0], b[0]], [a[1], b[1]], "-", color="#3b82f6", lw=1.6)
    for (a, b) in raw_lines:
        ax.plot(a[0], a[1], "o", color="#1e40af", ms=5)
        ax.plot(b[0], b[1], "o", color="#1e40af", ms=5)
    ax.set_title("(a) 원본 LINE 끝점\n(좌표에 미세 오차 — 점들이 따로 떠 있음)", fontsize=10)
    ax.set_xlim(-0.5, 7); ax.set_ylim(-0.5, 2.5); ax.set_aspect("equal")
    ax.grid(True, alpha=0.25)
    # (b) Snap 적용 (격자 정렬)
    ax = axes[1]
    grid = 0.2
    def snap(p): return (round(p[0]/grid)*grid, round(p[1]/grid)*grid)
    edges = set()
    nodes = set()
    for a, b in raw_lines:
        a2 = snap(a); b2 = snap(b)
        nodes.add(a2); nodes.add(b2)
        edges.add((min(a2, b2), max(a2, b2)))
    for (a, b) in edges:
        ax.plot([a[0], b[0]], [a[1], b[1]], "-", color="#0ea5e9", lw=1.8)
    for n in nodes:
        ax.plot(n[0], n[1], "o", color="#0c4a6e", ms=6)
    ax.set_title("(b) Snap 200mm 격자로 통합\n(가까운 끝점들이 한 노드로)", fontsize=10)
    ax.set_xlim(-0.5, 7); ax.set_ylim(-0.5, 2.5); ax.set_aspect("equal")
    ax.grid(True, alpha=0.25)
    # (c) Bridge — 끊어진 component 연결
    ax = axes[2]
    for (a, b) in edges:
        ax.plot([a[0], b[0]], [a[1], b[1]], "-", color="#0ea5e9", lw=1.8)
    # 가상 다리 — (3,0) ~ (5,0)
    ax.plot([3, 5], [0, 0], "--", color="#dc2626", lw=2.0, label="가상 다리")
    for n in nodes:
        ax.plot(n[0], n[1], "o", color="#0c4a6e", ms=6)
    ax.legend(loc="upper right", fontsize=9)
    ax.set_title("(c) 컴포넌트 가상 다리\n(분리된 component 들 연결)", fontsize=10)
    ax.set_xlim(-0.5, 7); ax.set_ylim(-0.5, 2.5); ax.set_aspect("equal")
    ax.grid(True, alpha=0.25)
    plt.tight_layout()
    path = FIG_DIR / "fig2_snap_bridge.png"
    plt.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close()
    return path


def fig_head_detection():
    """그림 3 — 헤드 인식 4 규칙."""
    fig, axes = plt.subplots(1, 4, figsize=(13, 3.5))
    titles_descs = [
        ("R1: Block 이름 매칭",
         "INSERT 블록 이름이\n알려진 head 블록\n(A$C39172136 등)\n→ 신뢰도 0.95"),
        ("R2: CIRCLE 시그니처",
         "HEAD 레이어의 원\n(반경 10~250mm)\n= 헤드 본체 마커\n→ 신뢰도 0.80"),
        ("R3+R5: HATCH 삼각형",
         "3 정점 채움 삼각형\n(드라이팬던트 마커)\n레이어 무관\n→ 신뢰도 0.72~0.75"),
        ("R4: 250mm 클러스터링",
         "여러 cue 가 한 자리에\n겹치면 한 헤드로 통합\n신뢰도 보너스\n→ 최대 0.99"),
    ]
    for ax, (title, desc) in zip(axes, titles_descs):
        ax.axis("off")
        ax.add_patch(patches.FancyBboxPatch((0.05, 0.7), 0.9, 0.25, boxstyle="round,pad=0.02",
                                             facecolor="#f97316", alpha=0.18, edgecolor="#f97316"))
        ax.text(0.5, 0.825, title, ha="center", va="center", fontsize=11, fontweight="bold")
        ax.text(0.5, 0.35, desc, ha="center", va="center", fontsize=9.5)
    fig.suptitle("Stage 2 — 헤드 인식: 4가지 신호 결합", fontsize=12, fontweight="bold")
    plt.tight_layout()
    path = FIG_DIR / "fig3_head_detection.png"
    plt.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close()
    return path


def fig_dijkstra_concept():
    """그림 4 — Dijkstra 가장 먼 30 헤드 선정."""
    fig, axes = plt.subplots(1, 2, figsize=(11, 5))
    # 토폴로지 좌측
    ax = axes[0]
    np.random.seed(7)
    src = (0, 0)
    heads = []
    edges_drawn = []
    # 트리 모양 가지
    main_branches = [(np.cos(a), np.sin(a)) for a in np.linspace(0, 2*np.pi, 5, endpoint=False)]
    for b in main_branches:
        bx, by = b[0] * 1.5, b[1] * 1.5
        edges_drawn.append((src, (bx, by)))
        for j in range(3):
            ang = np.arctan2(by, bx) + (j - 1) * 0.3
            hx = bx + np.cos(ang) * 1.8 + np.random.uniform(-0.2, 0.2)
            hy = by + np.sin(ang) * 1.8 + np.random.uniform(-0.2, 0.2)
            edges_drawn.append(((bx, by), (hx, hy)))
            heads.append((hx, hy))
    # 그리기
    for a, b in edges_drawn:
        ax.plot([a[0], b[0]], [a[1], b[1]], "-", color="#3b82f6", lw=1.5, alpha=0.6)
    for h in heads:
        ax.plot(h[0], h[1], "o", color="#ef4444", ms=8, mec="white", mew=1)
    ax.plot(src[0], src[1], "*", color="#a855f7", ms=22, mec="white", mew=1.5, label="알람밸브 (source)")
    ax.set_title("(a) 배관망 그래프\n각 노드까지 src→head 최단 경로 계산", fontsize=10)
    ax.legend(loc="upper right", fontsize=9)
    ax.set_xlim(-4, 4); ax.set_ylim(-4, 4); ax.set_aspect("equal")
    ax.axis("off")
    # 우측 — 가장 먼 N개 표시
    ax = axes[1]
    for a, b in edges_drawn:
        ax.plot([a[0], b[0]], [a[1], b[1]], "-", color="#cbd5e1", lw=1, alpha=0.5)
    # 거리별 정렬 후 상위 K
    dists = [(h, np.hypot(h[0], h[1])) for h in heads]
    dists.sort(key=lambda x: -x[1])
    K = 5
    worst = [h for h, _ in dists[:K]]
    # 경로 강조 (직선 근사 — 시각화용)
    for h in worst:
        # src → 가장 가까운 branch → head
        bx = (h[0] / np.hypot(h[0], h[1])) * 1.5
        by = (h[1] / np.hypot(h[0], h[1])) * 1.5
        ax.plot([src[0], bx], [src[1], by], "-", color="#fde047", lw=2.6)
        ax.plot([bx, h[0]], [by, h[1]], "-", color="#fde047", lw=2.6)
    for h in heads:
        c = "#ef4444" if h in worst else "#94a3b8"
        ms = 10 if h in worst else 6
        ax.plot(h[0], h[1], "o", color=c, ms=ms, mec="white", mew=1)
    ax.plot(src[0], src[1], "*", color="#a855f7", ms=22, mec="white", mew=1.5)
    ax.set_title(f"(b) 가장 먼 {K} 헤드 + 경로 (예시)\n빨간 점 = 선정된 헤드, 노란 선 = critical path",
                 fontsize=10)
    ax.set_xlim(-4, 4); ax.set_ylim(-4, 4); ax.set_aspect("equal")
    ax.axis("off")
    plt.tight_layout()
    path = FIG_DIR / "fig4_dijkstra.png"
    plt.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close()
    return path


def fig_user_edit_modes():
    """그림 5 — 사용자 편집 모드 3종."""
    fig, axes = plt.subplots(1, 3, figsize=(12, 4))
    np.random.seed(3)
    base_pts = np.random.uniform(0, 5, size=(10, 2))
    for ax, (title, desc, color) in zip(axes, [
        ("＋ 헤드 추가\n(클릭)", "캔버스 클릭 시\n200×200mm 박스 추가", "#16a34a"),
        ("− 헤드 삭제\n(클릭)", "헤드 근처 250mm 안\n클릭 시 삭제 마킹", "#dc2626"),
        ("▭ 영역 지정\n(드래그)", "드래그로 사각형\n안의 헤드만 사용", "#1e40af"),
    ]):
        ax.set_xlim(0, 5); ax.set_ylim(0, 5); ax.set_aspect("equal")
        ax.set_facecolor("#0f172a")
        # 기존 헤드 박스 (주황)
        for p in base_pts:
            ax.add_patch(patches.Rectangle((p[0]-0.15, p[1]-0.15), 0.3, 0.3,
                                            linewidth=1.2, edgecolor="#f97316", facecolor="none"))
        if "추가" in title:
            ax.add_patch(patches.Rectangle((3.7, 3.5), 0.35, 0.35,
                                            linewidth=2.5, edgecolor="#16a34a", facecolor="none"))
            ax.plot([3.7+0.175, 3.7+0.175], [3.5+0.08, 3.5+0.27], "-", color="#16a34a", lw=2)
            ax.plot([3.7+0.08, 3.7+0.27], [3.5+0.175, 3.5+0.175], "-", color="#16a34a", lw=2)
        elif "삭제" in title:
            # X 표시된 헤드
            p = base_pts[2]
            ax.add_patch(patches.Rectangle((p[0]-0.15, p[1]-0.15), 0.3, 0.3,
                                            linewidth=2.0, edgecolor="#dc2626", facecolor="none",
                                            linestyle="--"))
            ax.plot([p[0]-0.15, p[0]+0.15], [p[1]-0.15, p[1]+0.15], "-", color="#dc2626", lw=1.8)
            ax.plot([p[0]-0.15, p[0]+0.15], [p[1]+0.15, p[1]-0.15], "-", color="#dc2626", lw=1.8)
        else:  # 영역
            ax.add_patch(patches.Rectangle((1.2, 1.3), 2.5, 2.0,
                                            linewidth=2.0, edgecolor="#1e40af", facecolor="none",
                                            linestyle="--"))
        ax.text(2.5, 4.7, title, ha="center", color="white", fontsize=11, fontweight="bold")
        ax.text(2.5, 0.35, desc, ha="center", color="#cbd5e1", fontsize=9)
        ax.set_xticks([]); ax.set_yticks([])
    plt.tight_layout()
    path = FIG_DIR / "fig5_user_edit.png"
    plt.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close()
    return path


def fig_io_overview():
    """그림 6 — 입력 1개 → 출력 7개 흐름도."""
    fig, ax = plt.subplots(figsize=(11, 5))
    ax.set_xlim(0, 12); ax.set_ylim(0, 7); ax.axis("off")
    # 입력 박스
    ax.add_patch(patches.FancyBboxPatch((0.3, 2.8), 2.2, 1.4, boxstyle="round,pad=0.05",
                                         linewidth=1.6, edgecolor="#1e40af", facecolor="#dbeafe"))
    ax.text(1.4, 3.5, "DXF 도면\n(.dxf)", ha="center", va="center", fontsize=12, fontweight="bold")
    # 화살표
    ax.annotate("", xy=(3.0, 3.5), xytext=(2.6, 3.5),
                arrowprops=dict(arrowstyle="->", color="#475569", lw=2))
    # 모듈 박스
    ax.add_patch(patches.FancyBboxPatch((3.3, 1.8), 3.0, 3.4, boxstyle="round,pad=0.05",
                                         linewidth=2, edgecolor="#7c3aed",
                                         facecolor="#ede9fe"))
    ax.text(4.8, 4.6, "Remote 30 프로토타입", ha="center", fontsize=11, fontweight="bold", color="#5b21b6")
    ax.text(4.8, 3.7, "6-Stage 자동 처리\n+ 사용자 편집", ha="center", va="center", fontsize=10)
    ax.text(4.8, 2.5, "캔버스 실시간\n시각화", ha="center", va="center", fontsize=9.5, style="italic", color="#5b21b6")
    # 화살표
    ax.annotate("", xy=(6.9, 3.5), xytext=(6.4, 3.5),
                arrowprops=dict(arrowstyle="->", color="#475569", lw=2))
    # 출력 7개
    outputs = [
        ("XLSX\n(6 시트)", 6.5),
        ("SDF\n(PIPENET)", 5.5),
        ("CSV Nodes", 4.5),
        ("CSV Pipes", 3.5),
        ("CSV Nozzles", 2.5),
        ("CSV Fittings", 1.5),
        ("CSV Equipment", 0.5),
    ]
    for label, y in outputs:
        ax.add_patch(patches.FancyBboxPatch((7.1, y), 1.7, 0.8, boxstyle="round,pad=0.02",
                                             linewidth=1.2, edgecolor="#0f766e",
                                             facecolor="#d1fae5"))
        ax.text(7.95, y + 0.4, label, ha="center", va="center", fontsize=9.5, fontweight="bold")
    ax.text(11.0, 3.5, "사용자 다운로드", ha="center", va="center", fontsize=10, color="#0f766e",
            rotation=270)
    ax.text(6.0, 6.3, "입력 1개 → 출력 7개", ha="center", fontsize=13, fontweight="bold")
    plt.tight_layout()
    path = FIG_DIR / "fig6_io.png"
    plt.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close()
    return path


# ────────────────────────────────────────────────────────────────────────────
# Word 보고서 빌드
# ────────────────────────────────────────────────────────────────────────────


def set_korean_font(doc):
    """모든 스타일에 한글 폰트 적용."""
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        font = style.font
        font.name = "Malgun Gothic"
        rpr = style.element.rPr
        if rpr is not None:
            r = rpr.find(qn("w:rFonts"))
            if r is None:
                r = OxmlElement("w:rFonts")
                rpr.append(r)
            r.set(qn("w:eastAsia"), "Malgun Gothic")
            r.set(qn("w:ascii"), "Malgun Gothic")
            r.set(qn("w:hAnsi"), "Malgun Gothic")


def add_centered_image(doc, image_path, width_cm=15, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def build_doc(figs):
    doc = Document()
    set_korean_font(doc)

    # 표지
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Remote 30 프로토타입 모듈\n작동 원리 보고서")
    r.bold = True; r.font.size = Pt(22)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("CAD 도면(DXF) 한 장 → PIPENET 입력 파일까지 자동 변환\n— 일반인용 해설서 —")
    r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = info.add_run("작성: 2026.05  ·  대상 독자: 비전공자")
    r.italic = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    doc.add_page_break()

    # 0. 한눈에 보는 모듈
    doc.add_heading("0. 한눈에 보는 모듈", level=1)
    doc.add_paragraph(
        "이 모듈의 역할은 한 마디로 — “건축 설계용 CAD 도면(DXF) 한 장을 받으면, "
        "수리 계산 프로그램(PIPENET)이 바로 읽을 수 있는 입력 파일을 자동으로 만들어주는 것” 입니다."
    )
    doc.add_paragraph(
        "스프링클러 도면을 사람이 손으로 옮겨 입력하면 보통 수 시간 ~ 하루가 걸리는 작업을, "
        "이 모듈은 1 분 안에 끝냅니다. 그 과정에서 어떤 데이터를 어떻게 다루는지를 "
        "이 보고서가 차근차근 설명합니다."
    )
    add_centered_image(doc, figs["io"], width_cm=15,
                       caption="[그림 6] 입력 1개 (DXF) → 출력 7개 (XLSX, SDF, CSV 5종) 흐름도")

    # 1. 도면이 컴퓨터 안에서 어떻게 보이나
    doc.add_heading("1. 도면이 컴퓨터 안에서 어떻게 보이나", level=1)
    doc.add_paragraph(
        "사람이 도면을 볼 때는 “배관”, “헤드”, “벽” 같은 의미 단위로 인식합니다. "
        "그런데 컴퓨터에게 DXF 파일은 그저 수십만 개의 작은 도형(선/원/문자/블록)이 좌표 위에 놓인 목록일 뿐입니다."
    )
    p = doc.add_paragraph()
    p.add_run("예) 대명동201동 단위세대 도면 → 약 40,000 개의 entity가 들어 있고, "
              "그 중 의미 있는 배관 관련 객체는 약 2,500 개 입니다. ").italic = True
    doc.add_paragraph(
        "그래서 첫 번째 과제는 “수많은 도형 더미에서 배관·헤드만 골라내는 것” 이고, "
        "두 번째 과제는 “골라낸 도형들을 ‘배관망 = 노드와 길로 이루어진 그물망’ 으로 재해석하는 것” 입니다."
    )

    # 2. 전체 흐름
    doc.add_heading("2. 전체 흐름 — 6단계 자동 + 1번 사람 개입", level=1)
    add_centered_image(doc, figs["flow"], width_cm=16,
                       caption="[그림 1] 6-Stage 자동 파이프라인 흐름도 — Stage 3 후 사용자 편집 단계가 한 번 들어갑니다")
    doc.add_paragraph(
        "버튼 한 번 클릭으로 0→1→2→3 까지 자동 실행되고, 거기서 잠시 멈춥니다. "
        "사용자가 “헤드를 추가/삭제하거나 분석 영역을 지정” 한 뒤 “배관망 완성” 버튼을 누르면 "
        "4→5→6 이 자동 실행되며 최종 결과 파일이 만들어집니다."
    )

    # 3. Stage별 설명
    doc.add_heading("3. 6단계 자세히 보기", level=1)

    # Stage 0
    doc.add_heading("Stage 0 — DXF 파싱과 좌표 보정", level=2)
    doc.add_paragraph(
        "CAD 파일에는 “블록” 이라는 묶음이 있습니다. 예를 들어 ‘스프링클러 헤드 한 개’ 가 "
        "원 + 삼각형 + 글자 여러 도형의 묶음으로 정의되어 있고, 도면 안에서는 그 묶음이 "
        "다양한 위치 · 회전 · 크기로 반복 등장합니다."
    )
    doc.add_paragraph(
        "이 모듈은 그 블록들을 모두 풀어 헤쳐서 (explode), 도면 어디에 어떤 도형이 정확히 어디 있는지 "
        "한꺼번에 좌표로 정리합니다. 동시에 CAD 에서 “꺼짐(off)” 또는 “동결(frozen)” 처리된 레이어 "
        "(살수 반경 빨간 원처럼 사람 눈에는 안 보이는 것들) 는 미리 걸러냅니다."
    )

    # Stage 1
    doc.add_heading("Stage 1 — “배관망만” 추리기", level=2)
    doc.add_paragraph(
        "스프링클러 도면에는 건축선·치수선·소화기·콘센트 등 우리에게 필요 없는 도형이 90% 이상을 "
        "차지합니다. 레이어 이름을 단서로 “배관 / 헤드 / 텍스트” 만 통과시키고 나머지는 버립니다."
    )
    p = doc.add_paragraph()
    p.add_run("실측: 입력 42,440개 entity → Stage 1 후 2,512개로 축소 (94% 제거).").italic = True

    # Stage 2
    doc.add_heading("Stage 2 — 헤드 인식 (바운딩박스)", level=2)
    add_centered_image(doc, figs["heads"], width_cm=16,
                       caption="[그림 3] Stage 2 의 4가지 헤드 인식 규칙")
    doc.add_paragraph(
        "스프링클러 헤드는 도면에서 보통 [블록 + 원 + 삼각형 하나의 묶음] 으로 그려집니다. "
        "이 모듈은 4가지 단서를 결합해 “여기에 헤드가 있다” 라고 판단합니다:"
    )
    bullets = [
        "R1) 블록 이름이 우리가 아는 헤드 블록과 일치 (신뢰도 0.95)",
        "R2) 헤드 레이어의 원이 헤드 크기(반경 10~250mm) 범위 (신뢰도 0.80)",
        "R3 / R5) 삼각형 모양의 채움 도형 = 드라이팬던트 헤드 마커 (신뢰도 0.72~0.75)",
        "R4) 위 신호들이 250mm 안에 여러 개 겹치면 한 헤드로 합침 (신뢰도 +)",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")
    doc.add_paragraph(
        "이렇게 인식한 모든 헤드 위치는 화면에 주황색 박스로 표시됩니다. "
        "대명동 도면 기준 약 118개의 헤드 후보가 평균 신뢰도 0.93 으로 잡힙니다."
    )

    # Stage 3
    doc.add_heading("Stage 3 — 배관망을 “그래프”로 인식", level=2)
    add_centered_image(doc, figs["snap"], width_cm=16,
                       caption="[그림 2] (a) 원본 CAD LINE → (b) 200mm 격자에 스냅 → (c) 끊어진 곳을 가상 다리로 연결")
    doc.add_paragraph(
        "CAD 도면의 선들은 끝점이 살짝씩 어긋나 있어 (예: 0.5mm) 컴퓨터는 “이 선과 저 선이 한 점에서 만난다” 라고 "
        "쉽게 인식하지 못합니다. 이 단계에서는 다음 3 가지를 합니다:"
    )
    doc.add_paragraph("① Snap — 200mm 격자에 맞춰 가까운 끝점들을 “같은 점”으로 통합", style="List Bullet")
    doc.add_paragraph("② Bridge — 그래도 끊어진 배관 조각들을 점진적 거리(200·500·1,000·2,000mm)로 가상 다리 연결", style="List Bullet")
    doc.add_paragraph("③ Drop line — 헤드 위치를 가장 가까운 배관에 연결 (실제 도면에서 헤드 INSERT 마커가 배관에서 살짝 떨어져 있는 경우가 많음)", style="List Bullet")
    doc.add_paragraph(
        "결과적으로 약 1,190 개의 노드(꺾이는 점·분기점·헤드 위치)와 1,157 개의 길(edge), "
        "그리고 228 개의 분기점(차수≥3 인 “T 자 / 십자 갈림길”) 을 가진 그래프가 만들어집니다. "
        "이 그래프 위에서 다음 단계가 작동합니다."
    )

    # 사용자 편집
    doc.add_heading("👤 사용자 편집 단계 (자동 실행 일시 정지)", level=2)
    add_centered_image(doc, figs["edit"], width_cm=16,
                       caption="[그림 5] 3가지 편집 모드 — 헤드 추가 / 헤드 삭제 / 영역 지정")
    doc.add_paragraph(
        "Stage 3 까지 자동 진행 후 화면이 잠시 멈추고 “헤드 객체 수정” 패널이 활성화됩니다. "
        "여기서 사용자는:"
    )
    doc.add_paragraph("• 컴퓨터가 못 찾은 헤드를 캔버스 클릭으로 추가 (200×200mm 박스)", style="List Bullet")
    doc.add_paragraph("• 잘못 잡힌 헤드를 클릭으로 삭제", style="List Bullet")
    doc.add_paragraph("• 분석할 영역을 드래그로 지정 (영역 밖 헤드는 무시됨)", style="List Bullet")
    doc.add_paragraph("• 알람밸브 좌표를 직접 입력하거나 캔버스 클릭으로 지정", style="List Bullet")
    doc.add_paragraph(
        "편집이 끝나면 파란 [▶ 배관망 완성] 버튼으로 나머지 단계를 진행합니다."
    )

    # Stage 4
    doc.add_heading("Stage 4 — 가장 불리한 30개 헤드 + 경로 색출", level=2)
    add_centered_image(doc, figs["dijkstra"], width_cm=16,
                       caption="[그림 4] 알람밸브에서 각 헤드까지의 최단 경로 → 그 중 가장 먼 30개 + 합집합 경로")
    doc.add_paragraph(
        "수리 계산의 출발점인 “알람밸브(Alarm Valve, 가압수가 진입하는 자리)” 를 그래프에서 자동 식별합니다. "
        "그 후 다익스트라(Dijkstra) 알고리즘으로 알람밸브에서 모든 헤드까지의 “파이프를 따라 흐르는 거리” 를 계산합니다."
    )
    doc.add_paragraph(
        "가장 거리가 먼 30개 헤드를 뽑고, 각 헤드까지의 경로를 모두 합치면 “수압이 가장 떨어지는 임계 배관 망(critical subgraph)” 이 됩니다. "
        "수리 계산은 보통 이 가장 불리한 케이스 30개를 기준으로 안전 여부를 판단합니다."
    )

    # Stage 5
    doc.add_heading("Stage 5 — 입력 테이블 5종 + Meta 만들기", level=2)
    doc.add_paragraph(
        "선정된 헤드 30개와 거기까지의 배관망을 표로 정리합니다. PIPENET 이 요구하는 6 시트(엑셀) 구조 그대로:"
    )
    doc.add_paragraph("• Nodes — 노드 라벨, 좌표(mm), 높이(m), 입력/일반 구분", style="List Bullet")
    doc.add_paragraph("• Pipes — 파이프 라벨, 양 끝 노드, 호칭경(mm), 길이(m), C-factor", style="List Bullet")
    doc.add_paragraph("• Nozzles — 헤드(노즐) 라벨, 유량 (80 L/min, SP-HEAD)", style="List Bullet")
    doc.add_paragraph("• Fittings — 각 파이프 위의 부속(티/엘보) 개수", style="List Bullet")
    doc.add_paragraph("• Equipment — 알람밸브(A/V) · 후렉시블(FX) 등 등가 길이 부속", style="List Bullet")
    doc.add_paragraph("• Meta — 원본 파일명 · SDF 버전 · 생성 시각 등 메타데이터", style="List Bullet")
    doc.add_paragraph(
        "파이프의 관경(호칭경) 은 도면의 TEXT 레이어에 적힌 ‘25A / 32A / 50A / …’ 같은 표기를 "
        "주변 5m 안에서 자동 매칭해서 채워 넣습니다. "
        "후렉시블 길이는 ‘SP 후렉시블’ 레이어의 폴리라인 길이를 그대로 사용합니다."
    )

    # Stage 6
    doc.add_heading("Stage 6 — PIPENET SDF 파일로 출력", level=2)
    doc.add_paragraph(
        "마지막으로 PIPENET 이 직접 열 수 있는 .sdf 파일을 만듭니다. "
        "내부적으로는 참조 SDF 한 개를 “템플릿” 으로 사용해 도면 모양(아이소매트릭 표시 옵션 · 색상 · 그리드) 은 그대로 두고, "
        "Nodes / Links 부분만 우리 데이터로 교체합니다. 그래서 PIPENET 에서 정상적으로 열리고 도식이 깨지지 않습니다."
    )
    doc.add_paragraph(
        "최종적으로 7개의 다운로드 가능한 파일이 만들어집니다: "
        "XLSX(6 시트) · SDF · CSV 5종(Nodes/Pipes/Nozzles/Fittings/Equipment)."
    )

    # 4. 핵심 알고리즘
    doc.add_heading("4. 핵심 알고리즘 — 직관적 설명", level=1)
    doc.add_heading("4-1. Snap (격자 정렬)", level=2)
    doc.add_paragraph(
        "각 끝점 좌표를 200mm 격자에 “반올림” 합니다. "
        "예: (12,345.7mm , −9,876.2mm) → (12,400 , −9,800). "
        "이렇게 하면 0.5mm 어긋나 있던 두 끝점이 자동으로 같은 노드로 인식됩니다. "
        "200mm 라는 값은 “PIPENET 참조 도면의 최소 파이프 길이 320mm” 보다 짧기 때문에 "
        "잘못된 통합은 거의 일어나지 않습니다."
    )
    doc.add_heading("4-2. Component Bridge (가상 다리)", level=2)
    doc.add_paragraph(
        "Snap 만으로도 끊어진 부분이 남으면, 가장 큰 “섬” 에서 가까운 작은 섬을 차례로 연결합니다. "
        "200mm → 500mm → 1,000mm → 2,000mm 순서로 가까운 다리를 먼저 놓아 자연스러운 연결을 우선시합니다. "
        "이게 없으면 일부 헤드까지 거리 계산이 ∞ 가 되어 선정에서 누락됩니다."
    )
    doc.add_heading("4-3. Dijkstra (최단경로)", level=2)
    doc.add_paragraph(
        "수십 년 된 고전 알고리즘으로, “시작점에서 모든 다른 점까지의 가장 짧은 경로” 를 정확하게 계산합니다. "
        "이 모듈에서는 “알람밸브에서 모든 헤드까지의 배관을 따라가는 거리” 를 한 번에 구합니다. "
        "거리 = 파이프 길이의 합이며, 단위는 mm 입니다."
    )
    doc.add_heading("4-4. Collinear Merge (직선 흡수)", level=2)
    doc.add_paragraph(
        "최단경로 위의 노드 중 “직선 위에 있는 중간 점” 은 결과 표에 안 필요하므로 흡수해서 파이프 길이로 합칩니다. "
        "꺾이는 점(엘보)이거나 갈라지는 점(티)만 노드로 남깁니다. "
        "결과적으로 PIPENET 표의 노드 수가 합리적인 수준으로 줄어듭니다 (대명동 도면: 259 → 약 128개)."
    )

    # 5. 신뢰성과 한계
    doc.add_heading("5. 신뢰성과 한계", level=1)
    doc.add_paragraph(
        "이 모듈은 “한국 NFTC 103 기준의 일반적인 sprinkler 도면” 을 가정해 설계되었습니다. "
        "다음 항목에서 잘 작동합니다:"
    )
    doc.add_paragraph("• 레이어 이름이 ‘-소화(SP가지관)’ , ‘-소화(SP헤드하향)’, ‘SP 후렉시블’ 같은 표준 컨벤션을 따를 때", style="List Bullet")
    doc.add_paragraph("• 헤드가 INSERT 블록 또는 CIRCLE 마커로 그려져 있을 때", style="List Bullet")
    doc.add_paragraph("• 호칭경이 TEXT 로 ‘25A / 50A / 100A’ 형태로 표기되어 있을 때", style="List Bullet")
    doc.add_paragraph("이런 한계가 있습니다:")
    doc.add_paragraph("• 레이어 이름이 비표준이면 자동 카테고리 분류가 실패할 수 있어요 (편집 단계에서 사용자가 보완)", style="List Bullet")
    doc.add_paragraph("• 알람밸브 위치 자동 식별이 어려운 도면은 수동 좌표 입력이 필요해요", style="List Bullet")
    doc.add_paragraph("• 다층 도면의 경우 1 단위 도면 단위로 처리하는 것이 결과가 깨끗합니다", style="List Bullet")

    # 6. 결론
    doc.add_heading("6. 결론", level=1)
    doc.add_paragraph(
        "Remote 30 프로토타입은 “사람이 도면을 의미적으로 이해하는 과정” 을 6 단계로 분해하여 "
        "각 단계마다 합리적인 규칙·알고리즘을 적용한 자동화 도구입니다. "
        "복잡한 인공지능 학습 모델 없이도 “레이어 이름 + 도형 모양 + 그래프 알고리즘” 만으로 "
        "PIPENET 수리계산이 받아들일 수 있는 입력 파일을 만들 수 있다는 것을 보여줍니다."
    )
    doc.add_paragraph(
        "처리 시간은 13MB 도면 기준 약 33 초 (Stage 0 의 DXF 파싱이 대부분), "
        "사용자 편집 후 Stage 4~6 마무리는 0.3 초 정도로 거의 즉시 완료됩니다."
    )
    doc.add_paragraph(
        "향후 개선 여지로는 (1) 레이어 자동 분류의 일반화, "
        "(2) 알람밸브 위치 식별의 정밀화, "
        "(3) 비표준 헤드 블록의 학습 기반 인식, "
        "(4) 다층 도면 동시 처리 같은 항목이 있습니다."
    )

    # 부록
    doc.add_page_break()
    doc.add_heading("부록 A. 용어 정리", level=1)
    terms = [
        ("DXF", "AutoCAD 의 텍스트 기반 도면 교환 포맷. 도면의 모든 선/원/문자 등을 좌표 목록으로 저장"),
        ("PIPENET", "영국 Sunrise 시스템의 배관 수리계산 소프트웨어. .sdf 가 입력 포맷"),
        ("SDF", "PIPENET 의 입력 파일 (XML 기반)"),
        ("Entity", "DXF 안의 한 개 도형 (LINE / CIRCLE / TEXT / INSERT 등)"),
        ("INSERT", "도면 안에서 “블록 정의” 를 한 번 인용한 것 (헤드/티 같은 반복 부품)"),
        ("Layer", "도면의 “레이어” — 배관용 · 건축용 · 문자용 등 분류"),
        ("Snap", "끝점들을 같은 격자 위치로 반올림해 같은 노드로 통합"),
        ("Bridge", "분리된 배관 그래프 조각들을 가까운 끝점끼리 가상 edge 로 잇는 것"),
        ("Drop line", "헤드 위치 ↔ 가장 가까운 배관 노드를 잇는 짧은 가상 edge"),
        ("Dijkstra", "그래프에서 한 점에서 모든 점까지의 최단 거리를 계산하는 알고리즘"),
        ("Critical subgraph", "가장 불리한 N 헤드 각각의 최단경로를 합친 부분 그래프 — 수리계산의 핵심 대상"),
        ("Nozzle", "PIPENET 용어로 스프링클러 헤드 (1 nozzle = 1 head)"),
        ("A/V", "Alarm Valve — 알람밸브. 가압수의 진입점"),
        ("FX", "Flexible (후렉시블) — 신축성 있는 짧은 연결 호스 부속"),
        ("Equivalent length", "부속의 마찰손실을 직선 파이프 길이로 환산한 값 (m 단위)"),
        ("C-factor", "Hazen-Williams 식의 마찰계수 (KSD 3507 강관은 120)"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "용어"; hdr_cells[1].text = "설명"
    for term, desc in terms:
        row = table.add_row().cells
        row[0].text = term; row[1].text = desc

    return doc


def main():
    print("[1] 삽도 생성...")
    figs = {
        "flow": fig_pipeline_flow(),
        "snap": fig_snap_graph_concept(),
        "heads": fig_head_detection(),
        "dijkstra": fig_dijkstra_concept(),
        "edit": fig_user_edit_modes(),
        "io": fig_io_overview(),
    }
    for k, v in figs.items():
        print(f"  {k}: {v.name}  ({v.stat().st_size//1024} KB)")

    print("[2] 워드 문서 빌드...")
    doc = build_doc(figs)
    out_path = OUT_DIR / "Remote30_프로토타입_원리_보고서.docx"
    doc.save(out_path)
    print(f"\n[done] {out_path}  ({out_path.stat().st_size/1024:.1f} KB)")


if __name__ == "__main__":
    main()
