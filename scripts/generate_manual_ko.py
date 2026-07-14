# -*- coding: utf-8 -*-
from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def build_manual(docx_path: Path, image_path: Path | None = None) -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    style.font.size = Pt(10.5)

    doc.add_heading("PIPENET 수리계산 검증 프로그램 설명서", 0)
    meta = doc.add_paragraph()
    meta.add_run(
        f"문서 버전: V1.1\n작성일: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n대상 시스템: Flask 기반 웹 검증 서버"
    )
    doc.add_paragraph("※ 본 문서는 참조 설명서의 번호형 구성(개요-기준-화면-운영-장애대응)을 반영해 작성했습니다.")

    doc.add_heading("1. 프로그램 개요", level=1)
    doc.add_paragraph(
        "이 프로그램은 PIPENET 결과서(docx/pdf)와 SDF 파일(선택)을 입력받아, "
        "가이드라인 기준에 따라 자동 검증하고 결과를 웹 화면 및 엑셀 파일로 제공합니다."
    )
    for item in [
        "법적/기술 기준 위반 항목 자동 탐지",
        "공학적 최적화 후보(마찰손실 급증 구간) 식별",
        "경제성 검토 후보(과설계 가능 구간) 식별",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("2. 입력 파일과 처리 범위", level=1)
    doc.add_paragraph("2.1 입력 파일")
    doc.add_paragraph("결과서 파일: .docx 또는 .pdf (필수)", style="List Bullet")
    doc.add_paragraph("SDF 파일: .sdf (선택)", style="List Bullet")
    doc.add_paragraph("2.2 결과서에서 읽는 주요 섹션")
    for item in [
        "FLOW IN PIPES",
        "FLOW THROUGH NOZZLES",
        "SPECIAL EQUIPMENT",
        "FLOW THROUGH ELASTOMERIC VALVES",
        "PIPE CONFIGURATION(길이 정보)",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph("2.3 SDF에서 읽는 주요 정보")
    for item in [
        "Title(메인/존)",
        "Node, Pipe, Nozzle, Equipment 구조",
        "좌표 기반 아이소매트릭 배관망 정보",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3. 검증 로직(핵심 기준)", level=1)
    doc.add_paragraph("3.1 배관 유속 기준")
    for item in [
        "50A 이하: 6.0 m/s 이하",
        "50A 초과: 10.0 m/s 이하",
        "기준 초과 시 부적합(FAIL) 처리",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph("3.2 헤드(노즐) 기준")
    for item in [
        "최소 유량: 80 L/min 이상",
        "허용 압력 범위: 1.0 ~ 12.0 kg/cm²",
        "미달 또는 범위 이탈 시 부적합(FAIL) 처리",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph("3.3 특수설비/밸브 기준")
    for item in [
        "FX 등가길이: 13 ~ 21 m",
        "A/V 등가길이: 12.9 m (±0.1)",
        "P/V 등가길이: 10.1 m (±0.1)",
        "기준 불일치 시 부적합(FAIL), 일부 누락은 확인 필요(WARNING) 처리",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph("3.4 설계 최적화 로직")
    for item in [
        "공학 최적화 후보: 단위 길이당 마찰손실 급증 구간 탐지",
        "경제성 후보: 저유속 과설계 구간 및 고비용 밸브 규격 후보 탐지",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("4. 화면 구성 및 사용 방법", level=1)
    doc.add_paragraph("4.1 기본 사용 절차")
    for item in [
        "결과서 파일(docx/pdf) 업로드",
        "SDF 파일(선택) 업로드",
        "검증 실행 버튼 클릭",
        "좌측 메뉴에서 필요한 화면만 선택하여 확인",
    ]:
        doc.add_paragraph(item, style="List Number")
    doc.add_paragraph("4.2 좌측 메뉴 구성")
    for item in [
        "1. 검증결과",
        "2. 설계 최적화 가이드",
        "3. 결과 데이터 테이블",
        "4. 검진 통계",
        "5. 상세리포트",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph("4.3 결과 데이터 테이블 표시 규칙(색상)")
    for item in [
        "빨강: 기준 위반",
        "노랑: 확인 필요",
        "파랑: 공학 최적화 후보",
        "초록: 경제성 검토 후보",
        "필터링된 셀 클릭 시 판정 로직 설명 팝업 표시",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph("4.4 검진 통계 화면 배치")
    for item in [
        "상단 좌측: Pipe Velocity Check 그래프",
        "상단 우측: Nozzle Pressure-Flow 그래프",
        "하단 좌측: 결과서(DOCX/PDF) 통계 표",
        "하단 우측: SDF 통계 표",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("5. 출력 결과(웹 + 엑셀)", level=1)
    doc.add_paragraph("5.1 웹 출력")
    for item in [
        "검증 요약(pass/fail/warning)",
        "검증 상세 문장(항목별)",
        "테이블(배관/헤드/특수설비/감압밸브)",
        "SDF 아이소매트릭 배관망(줌/이동 가능)",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph("5.2 엑셀 출력")
    for item in [
        "4개 시트 분리: 배관, 헤드, 특수설비, 감압밸브",
        "행 색상으로 판정 상태 표시(빨강/노랑/파랑/초록)",
        "테두리/헤더/열 너비 자동 스타일 적용",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("6. 서버 실행 및 접속 안내", level=1)
    for item in [
        "실행: python run_server.py",
        "기본 접속: http://127.0.0.1:5050",
        "외부 접속(동일 네트워크): http://<서버PC_IP>:5050",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph("※ 방화벽에서 5050 포트 허용이 필요할 수 있습니다.")

    doc.add_heading("7. 장애 대응 가이드", level=1)
    for item in [
        '증상: "검증 중입니다..."에서 멈춤 -> 서버 로그(server.err.log) 확인',
        "증상: 그래프 글자 깨짐 -> 서버 재시작 후 재검증 실행",
        "증상: 화면 반영 안 됨 -> 브라우저 Ctrl+F5 강력 새로고침",
        "증상: 파일 파싱 실패 -> 파일 형식(docx/pdf/sdf) 및 섹션 구조 점검",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("8. 코드 구조(참고)", level=1)
    for item in [
        "대조 서버.py: Flask API, 시각화, SDF 그래프, 엑셀 다운로드",
        "pipenet_validator.py: 파싱/검증/통계/테이블 데이터 생성 핵심 로직",
        "templates/index.html: 화면 골격(UI)",
        "static/app.js: 인터랙션/렌더링/필터/메뉴 전환",
        "static/styles.css: 디자인/레이아웃/반응형 스타일",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("9. 화면 예시", level=1)
    doc.add_paragraph("아래 이미지는 검진 통계/시각화 영역 예시입니다.")
    if image_path and image_path.exists():
        doc.add_picture(str(image_path), width=Inches(6.6))
    else:
        doc.add_paragraph("(화면 이미지 파일을 찾지 못해 본문에 삽입하지 못했습니다.)")

    doc.add_heading("10. 향후 확장 제안", level=1)
    for item in [
        "층/존별 비교 대시보드 추가",
        "PDF 표 인식 정확도 보정(좌표 기반 파서 병행)",
        "검증 기준 버전 관리(V4, V5 등)와 선택형 실행",
        "리포트 자동 생성(요약+근거 테이블+시정조치안)",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(docx_path)


def main() -> None:
    desktop = Path.home() / "Desktop"
    out_docx = desktop / "PIPENET_수리계산_검증프로그램_설명서.docx"
    image = Path(r"C:\Users\admin\AppData\Local\Temp\ai-chat-attachment-12243213516581316316.png")
    build_manual(out_docx, image if image.exists() else None)
    print(str(out_docx))


if __name__ == "__main__":
    main()
